# Path: Swiss_bank_agent/backend/services/rag_service.py

import logging
import os
import hashlib
import json
import re
import time
import threading
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path
from enum import Enum
import argparse
import sys

# Document processing
import PyPDF2
from docx import Document

# Vector database and embeddings
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer


# BM25 for hybrid search
from rank_bm25 import BM25Okapi
import numpy as np
import torch

# AI for contextual enhancement
import anthropic
from anthropic.types import TextBlock

# Configuration
from dotenv import load_dotenv

import warnings
os.environ['ANONYMIZED_TELEMETRY'] = 'False'
warnings.filterwarnings("ignore", category=UserWarning, module="chromadb")
logging.getLogger('chromadb.telemetry.product.posthog').setLevel(logging.CRITICAL)
logging.getLogger('chromadb.telemetry').setLevel(logging.CRITICAL)
logging.getLogger('chromadb').setLevel(logging.WARNING)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)



class DuplicateHandlingMode(Enum):
    """Enum for different duplicate handling strategies"""
    SKIP = "skip"
    OVERWRITE = "overwrite"
    VERSION = "version"
    TIMESTAMP = "timestamp"

class ChunkingStrategy(Enum):
    """Enum for different text chunking strategies"""
    FIXED_SIZE = "fixed_size"
    SEMANTIC = "semantic"
    SENTENCE = "sentence"
    PARAGRAPH = "paragraph"

class RetrievalStrategy(Enum):
    """Enum for different retrieval strategies"""
    VECTOR_ONLY = "vector_only"
    BM25_ONLY = "bm25_only"
    HYBRID = "hybrid"
    CONTEXTUAL_HYBRID = "contextual_hybrid"  

class AnthropicContextualRAGService:
    """
    Advanced RAG Service implementing Anthropic's Contextual Retrieval Architecture
    
    Features:
    - Contextual Embeddings (50-100 tokens of context per chunk)
    - Contextual BM25 (keyword search with context)
    - Hybrid Search (semantic + BM25 + reranking)
    - Advanced document processing with metadata preservation
    - Sophisticated chunking strategies
    - Performance optimization with caching
    """
    
    def __init__(self, 
                 documents_directory: Optional[str] = None,
                 chroma_db_path: str = "./chroma_db", 
                 collection_name: str = "contextual_documents",
                 duplicate_mode: DuplicateHandlingMode = DuplicateHandlingMode.SKIP,
                 chunking_strategy: ChunkingStrategy = ChunkingStrategy.SEMANTIC,
                 retrieval_strategy: RetrievalStrategy = RetrievalStrategy.CONTEXTUAL_HYBRID,
                 embedding_model: str = "infgrad/stella_en_1.5B_v5",
                 quiet_mode: bool = False):
        """
        Initialize Advanced Contextual RAG Service
        
        Args:
            documents_directory: Path to documents folder 
            chroma_db_path: ChromaDB storage path
            collection_name: Collection name for vector storage
            duplicate_mode: How to handle duplicate/modified files
            chunking_strategy: Text chunking approach
            retrieval_strategy: Retrieval method (vector, BM25, hybrid, contextual_hybrid)
            embedding_model: Sentence transformer model for embeddings
            quiet_mode: If True, suppress initialization logging
        """
        # Setup paths
        self._setup_paths(documents_directory, chroma_db_path)
        
        # Service configuration
        self.collection_name = collection_name
        self.duplicate_mode = duplicate_mode
        self.chunking_strategy = chunking_strategy
        self.retrieval_strategy = retrieval_strategy
        self.embedding_model_name = embedding_model
        
        # Initialize components
        self.client = None
        self.collection = None
        self.embedding_model = None
        self.claude_client = None
        
        # BM25 components for hybrid search
        self.bm25_index = None
        self.bm25_documents = []
        self.bm25_metadata = []
        
        # Document processing settings
        self.chunk_size = 1000  
        self.chunk_overlap = 200
        self.context_tokens = 75  
        self.max_chunk_size = 2000
        self.min_chunk_size = 100
        
        # Supported file types
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
        
        # Processing tracking
        self.processed_files_log = {}
        self.file_versions = {}
        self.processing_stats = {}
        self.contextual_cache = {}  

        # Add rate limiting attributes
        self.api_call_delay = 1.5  
        self.last_api_call_time = 0
        self.api_call_lock = threading.Lock()


        # Log files
        self.processing_log_file = self.documents_dir / "processing_log.json"
        self.versions_log_file = self.documents_dir / "versions_log.json"
        self.stats_log_file = self.documents_dir / "processing_stats.json"
        self.bm25_index_file = self.documents_dir / "bm25_index.json"
        self.contextual_cache_file = self.documents_dir / "contextual_cache.json"

        # Set quiet mode first
        self.quiet_mode = quiet_mode

        # Configure logging level based on quiet mode
        if quiet_mode:
            # Temporarily suppress INFO level logs during initialization
            current_level = logger.level
            logger.setLevel(logging.WARNING)
            # Also suppress sentence_transformers logging
            logging.getLogger('sentence_transformers').setLevel(logging.WARNING)
            logging.getLogger('chromadb').setLevel(logging.ERROR)
        
        # Setup paths
        self._setup_paths(documents_directory, chroma_db_path)
        
        # Service configuration
        self.collection_name = collection_name
        self.duplicate_mode = duplicate_mode
        self.chunking_strategy = chunking_strategy
        self.retrieval_strategy = retrieval_strategy
        self.embedding_model_name = embedding_model

        # Initialize services
        self._initialize_claude()
        self._initialize_embedding_model()
        self._initialize_chroma()
        self._load_processing_logs()
        self._load_bm25_index()

        if quiet_mode:
            logger.setLevel(current_level)
            logging.getLogger('sentence_transformers').setLevel(logging.INFO)
            logging.getLogger('chromadb').setLevel(logging.INFO)

    def _setup_paths(self, documents_directory: Optional[str], chroma_db_path: str):
        """Setup file paths relative to current module"""
        current_file = Path(__file__)
        backend_dir = current_file.parent.parent
        
        # Documents directory setup
        if documents_directory:
            self.documents_dir = Path(documents_directory)
        else:
            project_root = backend_dir.parent
            self.documents_dir = project_root / "data" / "documents"
        
        # Vector database path setup
        if chroma_db_path.startswith('./'):
            self.chroma_db_path = backend_dir / chroma_db_path[2:]
        else:
            self.chroma_db_path = Path(chroma_db_path)
        
        # Create directories
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        self.chroma_db_path.parent.mkdir(parents=True, exist_ok=True)
        
    
    def _initialize_claude(self):
        """Initialize Claude API client for RAG service with dedicated API key"""
        try:
            backend_dir = Path(__file__).parent.parent
            env_path = backend_dir / '.env'
            load_dotenv(env_path)
            
            # CHANGE: Use RAG-specific API key with fallback
            api_key = os.getenv("RAG_API_KEY") or os.getenv("ANTHROPIC_API_KEY")
            
            if not api_key:
                if not getattr(self, 'quiet_mode', False):
                    logger.warning("Ã¢Å¡ Ã¯Â¸Â RAG_API_KEY environment variable not found - contextual enhancement disabled")
                self.claude_client = None
                return
            
            self.claude_client = anthropic.Anthropic(api_key=api_key)
            if not getattr(self, 'quiet_mode', False):
                logger.info("Ã¢Å“â€¦ Claude API client initialized for RAG service with dedicated key")
            
        except Exception as e:
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Failed to initialize Claude API for RAG service: {e}")
            self.claude_client = None

    def query_documents_universal(self, query: str, top_k: int = 8, 
                            metadata_filters: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Universal query method that works with any metadata structure"""
        
        # Check if collection is properly initialized
        if not self.collection:
            return {
                "success": False,
                "error": "ChromaDB collection not initialized",
                "documents": []
            }
        
        try:
            if metadata_filters:
                # Query with filters
                results = self.collection.query(
                    query_texts=[query],
                    n_results=top_k * 2,  # Get more for better filtering
                    where=metadata_filters,
                    include=["documents", "metadatas", "distances"]
                )
                
                # If filtered results are insufficient, expand search
                if not results["documents"] or len(results["documents"][0]) < top_k // 2:
                    logger.info("Filtered results insufficient, expanding search")
                    results = self.collection.query(
                        query_texts=[query],
                        n_results=top_k,
                        include=["documents", "metadatas", "distances"]
                    )
            else:
                # Standard contextual hybrid query - use existing method
                return self._query_contextual_hybrid(query, top_k)
            
            # Format and return results
            results_dict = {
                "documents": results.get("documents", []),
                "metadatas": results.get("metadatas", []),
                "distances": results.get("distances", []),
                "ids": results.get("ids", [])
            }
            
            formatted_results = self._format_query_results(results_dict, query, "universal_query")
            
            # Add AI-powered direct answer extraction
            if formatted_results.get("success") and formatted_results.get("documents"):
                direct_answer = self.extract_direct_answer(query, formatted_results["documents"])
                formatted_results["direct_answer"] = direct_answer
            
            return formatted_results
            
        except Exception as e:
            logger.error(f"Universal query failed: {e}")
            # Fallback to existing contextual hybrid method
            return self._query_contextual_hybrid(query, top_k)
    
    def _get_model_cache_path(self) -> Optional[Path]:
        """Get the cache path for the embedding model"""
        try:
            cache_root = Path.home() / ".cache" / "huggingface" / "hub"
            model_cache_name = f"models--{self.embedding_model_name.replace('/', '--')}"
            return cache_root / model_cache_name
        except Exception as e:
            logger.warning(f"Ã¢Å¡ Ã¯Â¸Â Could not determine cache path: {e}")
            return None

    def _check_model_cache(self, cache_path: Optional[Path]) -> bool:
        """Check if model is already cached locally"""
        try:
            if not cache_path or not cache_path.exists():
                return False
            
            # Look for essential model files
            config_files = list(cache_path.rglob("config.json"))
            model_files = (list(cache_path.rglob("*.bin")) + 
                        list(cache_path.rglob("*.safetensors")) + 
                        list(cache_path.rglob("pytorch_model.bin")) +
                        list(cache_path.rglob("model.safetensors")))
            tokenizer_files = list(cache_path.rglob("tokenizer.json"))
            
            # Must have at least config and model files
            has_config = len(config_files) > 0
            has_model = len(model_files) > 0
            
            if has_config and has_model:
                total_size = sum(f.stat().st_size for f in model_files if f.exists())
                size_mb = total_size / (1024 * 1024)
                
                if not getattr(self, 'quiet_mode', False):
                    logger.info(f"Ã°Å¸â€œâ€¹ Found cached model files ({size_mb:.1f} MB)")
                return True
            
            return False
            
        except Exception as e:
            logger.warning(f"Ã¢Å¡ Ã¯Â¸Â Error checking model cache: {e}")
            return False

    def _clear_model_cache(self) -> bool:
        """Clear the model cache"""
        try:
            cache_path = self._get_model_cache_path()
            if cache_path and cache_path.exists():
                import shutil
                shutil.rmtree(cache_path)
                logger.info(f"Ã°Å¸â€”â€˜Ã¯Â¸Â Cleared model cache")
                return True
            else:
                logger.info("Ã°Å¸â€œÂ No cache found to clear")
                return False
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error clearing cache: {e}")
            return False

    def _initialize_embedding_model(self):
        """Initialize sentence transformer model with better cache checking"""
        try:
            if not getattr(self, 'quiet_mode', False):
                logger.info(f"Ã°Å¸Â¤â€“ Initializing embedding model: {self.embedding_model_name}")
            
            # Load cache directories from .env
            backend_dir = Path(__file__).parent.parent
            env_path = backend_dir / '.env'
            load_dotenv(env_path)

            # Set cache directories
            cache_dir = os.getenv('HF_HOME', str(Path.home() / ".cache" / "huggingface"))
            
            if not getattr(self, 'quiet_mode', False):
                logger.info(f"Ã°Å¸â€œÂ Using cache directory: {cache_dir}")

            # Check if model is properly cached - but don't log conflicting info
            cache_path = Path(cache_dir) / "hub" / f"models--{self.embedding_model_name.replace('/', '--')}"
            is_cached = self._verify_model_cache(cache_path)
            
            # Only log cache status if we're confident about it
            if not getattr(self, 'quiet_mode', False) and is_cached:
                logger.info(f"Ã°Å¸â€œÂ¦ Model found in local cache")

            device = 'cuda' if torch.cuda.is_available() else 'cpu'
            
            # Initialize with explicit cache folder
            self.embedding_model = SentenceTransformer(
                self.embedding_model_name, 
                device=device,
                cache_folder=cache_dir
            )
                
            if not getattr(self, 'quiet_mode', False):
                logger.info(f"Ã¢Å“â€¦ Embedding model loaded successfully")
            
        except Exception as e:
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Failed to load embedding model: {e}")
            try:
                device = 'cuda' if torch.cuda.is_available() else 'cpu'
                cache_dir = os.getenv('HF_HOME', str(Path.home() / ".cache" / "huggingface"))
                
                self.embedding_model = SentenceTransformer(
                    'infgrad/stella_en_1.5B_v5', 
                    device=device, 
                    cache_folder=cache_dir
                )
                
                if not getattr(self, 'quiet_mode', False):
                    logger.info("Ã¢Å“â€¦ Loaded fallback embedding model")
            except Exception as e2:
                if not getattr(self, 'quiet_mode', False):
                    logger.error(f"Ã¢ÂÅ’ Failed to load fallback model: {e2}")
                self.embedding_model = None

    def _verify_model_cache(self, cache_path: Path) -> bool:
        """Simplified cache verification"""
        try:
            if not cache_path.exists():
                return False
            
            # Just check if any model files exist
            model_files = (list(cache_path.rglob("*.bin")) + 
                        list(cache_path.rglob("*.safetensors")))
            
            return len(model_files) > 0 and any(f.stat().st_size > 1024*1024 for f in model_files)
            
        except Exception:
            return False
        
    def _initialize_chroma(self):
        """Initialize ChromaDB client and collection"""
        try:
            self.client = chromadb.PersistentClient(
                path=str(self.chroma_db_path),
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            
            try:
                self.collection = self.client.get_collection(name=self.collection_name)
                if not getattr(self, 'quiet_mode', False):
                    logger.info(f"Ã¢Å“â€¦ Connected to existing ChromaDB collection: {self.collection_name}")
            except Exception:
                self.collection = self.client.create_collection(
                    name=self.collection_name,
                    metadata={"description": "Anthropic-style contextual document collection"}
                )
                if not getattr(self, 'quiet_mode', False):
                    logger.info(f"Ã¢Å“â€¦ Created new ChromaDB collection: {self.collection_name}")
                
        except Exception as e:
            self.collection = None
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Failed to initialize ChromaDB: {e}")
            raise

    def _load_processing_logs(self):
        """Load processing, version, and statistics logs"""
        # Load processed files log
        try:
            if self.processing_log_file.exists():
                with open(self.processing_log_file, 'r') as f:
                    self.processed_files_log = json.load(f)
                if not getattr(self, 'quiet_mode', False):
                    logger.info(f"Ã°Å¸â€œâ€ž Loaded processing log with {len(self.processed_files_log)} entries")
            else:
                self.processed_files_log = {}
        except Exception as e:
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Error loading processing log: {e}")
            self.processed_files_log = {}
        
        # Load file versions log
        try:
            if self.versions_log_file.exists():
                with open(self.versions_log_file, 'r') as f:
                    self.file_versions = json.load(f)
                if not getattr(self, 'quiet_mode', False):
                    logger.info(f"Ã°Å¸â€œÅ¡ Loaded versions log with {len(self.file_versions)} file groups")
            else:
                self.file_versions = {}
        except Exception as e:
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Error loading versions log: {e}")
            self.file_versions = {}
            
        # Load processing statistics
        try:
            if self.stats_log_file.exists():
                with open(self.stats_log_file, 'r') as f:
                    self.processing_stats = json.load(f)
            else:
                self.processing_stats = {}
        except Exception as e:
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Error loading stats log: {e}")
            self.processing_stats = {}
            
        # Load contextual cache
        try:
            if self.contextual_cache_file.exists():
                with open(self.contextual_cache_file, 'r') as f:
                    self.contextual_cache = json.load(f)
            else:
                self.contextual_cache = {}
        except Exception as e:
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Error loading contextual cache: {e}")
            self.contextual_cache = {}

    def _load_bm25_index(self):
        """Load BM25 index for hybrid search"""
        try:
            if self.bm25_index_file.exists():
                with open(self.bm25_index_file, 'r') as f:
                    bm25_data = json.load(f)
                    self.bm25_documents = bm25_data.get("documents", [])
                    self.bm25_metadata = bm25_data.get("metadata", [])
                    
                if self.bm25_documents:
                    # Tokenize documents for BM25
                    tokenized_docs = [doc.lower().split() for doc in self.bm25_documents]
                    self.bm25_index = BM25Okapi(tokenized_docs)
                    if not getattr(self, 'quiet_mode', False):
                        logger.info(f"Ã°Å¸â€œÅ  Loaded BM25 index with {len(self.bm25_documents)} documents")
        except Exception as e:
            if not getattr(self, 'quiet_mode', False):
                logger.error(f"Ã¢ÂÅ’ Error loading BM25 index: {e}")
            self.bm25_index = None
            self.bm25_documents = []
            self.bm25_metadata = []
    
    def _save_processing_logs(self):
        """Save all processing logs"""
        try:
            # Save processed files log
            with open(self.processing_log_file, 'w') as f:
                json.dump(self.processed_files_log, f, indent=2)
            
            # Save file versions log
            with open(self.versions_log_file, 'w') as f:
                json.dump(self.file_versions, f, indent=2)
                
            # Save processing statistics
            with open(self.stats_log_file, 'w') as f:
                json.dump(self.processing_stats, f, indent=2)
                
            # Save BM25 index
            with open(self.bm25_index_file, 'w') as f:
                json.dump({
                    "documents": self.bm25_documents,
                    "metadata": self.bm25_metadata
                }, f, indent=2)
                
            # Save contextual cache
            with open(self.contextual_cache_file, 'w') as f:
                json.dump(self.contextual_cache, f, indent=2)
                
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error saving logs: {e}")
    
    def _rate_limited_api_call(self, api_call_func, operation_type="general", max_retries=3, *args, **kwargs):
        """Simplified API call with exponential backoff"""
        with self.api_call_lock:
            for attempt in range(max_retries):
                current_time = time.time()
                time_since_last_call = current_time - self.last_api_call_time
                
                if time_since_last_call < self.api_call_delay:
                    sleep_time = self.api_call_delay - time_since_last_call
                    time.sleep(sleep_time)
                
                try:
                    response = api_call_func(*args, **kwargs)
                    self.last_api_call_time = time.time()
                    return response
                        
                except Exception as e:
                    error_msg = str(e)
                    
                    # Check if it's a 500 error
                    if "500" in error_msg and attempt < max_retries - 1:
                        backoff_time = (2 ** attempt) + self.api_call_delay
                        if not getattr(self, 'quiet_mode', False):
                            logger.warning(f"API 500 error, retrying in {backoff_time:.2f}s (attempt {attempt + 1}/{max_retries})")
                        time.sleep(backoff_time)
                        continue
                    else:
                        if not getattr(self, 'quiet_mode', False):
                            logger.warning(f"RAG API call failed: {e}")
                        self.last_api_call_time = time.time()
                        return None
            
            return None
    
    def _detect_structured_content_with_ai(self, chunk: str, document_metadata: Dict[str, Any]) -> bool:
        """Use AI to detect if chunk contains structured information"""
        if not self.claude_client:
            return False
        
        prompt = f"""Analyze this text chunk and determine if it contains structured information that should be preserved exactly as-is.

    Text: "{chunk}"
    Document: {document_metadata.get('source_file', 'Unknown')}

    Structured information includes:
    - Contact details (names, emails, phones, addresses)
    - Financial figures (amounts, percentages, dates)
    - Role assignments (who owns/manages what)
    - Policy requirements (specific rules, procedures)
    - Contract terms (clauses, conditions, obligations)
    - Technical specifications (APIs, configurations)

    Respond with just: TRUE or FALSE"""

        try:
            def make_claude_call():
                if self.claude_client is None:
                    raise ValueError("Claude client not initialized")
                return self.claude_client.messages.create(
                    model="claude-3-5-haiku-20241022",
                    max_tokens=10,
                    temperature=0,
                    messages=[{"role": "user", "content": prompt}]
                )
            
            response = self._rate_limited_api_call(make_claude_call)
            
            if response is not None:
                # Extract text from response
                response_text = ""
                for block in response.content:
                    if hasattr(block, 'text'):
                        response_text += block.text
                    elif isinstance(block, str):
                        response_text += block
                
                return "TRUE" in response_text.upper()
        except Exception as e:
            logger.warning(f"Ã¢Å¡ Ã¯Â¸Â AI content detection failed: {e}")
            pass
        
        return False
    
    def _generate_contextual_information(self, chunk: str, document_content: str, 
                                   document_metadata: Dict[str, Any]) -> str:
        """
        Generate contextual information for a chunk using Anthropic's approach
        WITH RATE LIMITING and improved content preservation
        """
        # Create cache key
        chunk_hash = hashlib.md5(chunk.encode()).hexdigest()[:16]
        doc_hash = hashlib.md5(document_content.encode()).hexdigest()[:16]
        cache_key = f"{doc_hash}_{chunk_hash}"
        
        # Check cache first
        if cache_key in self.contextual_cache:
            cached_context = self.contextual_cache[cache_key]
            return f"{cached_context}\n\n{chunk}"
        
        # Check if chunk contains structured information that should be preserved
        has_structured_info = self._detect_structured_content_with_ai(chunk, document_metadata)
        
        # If chunk has structured info, use minimal context to preserve it
        if has_structured_info:
            simple_context = f"From {document_metadata.get('source_file', 'document')}"
            self.contextual_cache[cache_key] = simple_context
            return f"{simple_context}\n\n{chunk}"
        
        # Generate context using Claude if available
        if self.claude_client:
            try:
                prompt = f"""Analyze this document chunk and provide brief context (30-50 tokens) for search retrieval in a corporate banking/fintech environment.

    Document: {document_metadata.get('source_file', 'Unknown')}
    Chunk: "{chunk}"

    Context should help identify:
    1. Document type (policy, procedure, contract, technical spec, etc.)
    2. Business domain (lending, payments, compliance, operations, etc.)
    3. Key entities or relationships mentioned
    4. Where this fits in the business process

    Keep it concise and business-focused.
    Respond with just the context, nothing else."""

                # Use rate-limited API call
                def make_claude_call():
                    if self.claude_client is None:
                        raise ValueError("Claude client not initialized")
                    return self.claude_client.messages.create(
                        model="claude-3-5-haiku-20241022",
                        max_tokens=100,
                        temperature=0.1,
                        messages=[{"role": "user", "content": prompt}]
                    )

                # CHANGE: Pass operation type for cost tracking
                message = self._rate_limited_api_call(make_claude_call, "contextual_enhancement")
                
                if message is None:
                    # Fallback if API call failed
                    return self._generate_fallback_context(chunk, document_metadata)
                
                context = ""
                for block in message.content:
                    if isinstance(block, TextBlock):
                        context += block.text
                    elif isinstance(block, str):
                        context += block
                
                context = context.strip()
                
                # Cache the result
                self.contextual_cache[cache_key] = context
                
                return f"{context}\n\n{chunk}"
                
            except Exception as e:
                logger.warning(f"Ã¢Å¡ Ã¯Â¸Â Failed to generate context with Claude: {e}")
                self.api_errors_count += 1
        
        # Fallback: Generate simple context without Claude
        return self._generate_fallback_context(chunk, document_metadata)

    def _generate_fallback_context(self, chunk: str, document_metadata: Dict[str, Any]) -> str:
        """Generate fallback context when Claude API is not available"""
        source_file = document_metadata.get('source_file', 'Unknown')
        file_type = document_metadata.get('file_type', 'document')
        
        fallback_context = f"This chunk is from {source_file}, a {file_type[1:]} document."
        
        # Cache fallback
        chunk_hash = hashlib.md5(chunk.encode()).hexdigest()[:16]
        doc_hash = hashlib.md5(f"{source_file}_{file_type}".encode()).hexdigest()[:16]
        cache_key = f"{doc_hash}_{chunk_hash}"
        self.contextual_cache[cache_key] = fallback_context
        
        return f"{fallback_context}\n\n{chunk}"
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file for duplicate detection"""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error calculating hash for {file_path}: {e}")
            return ""
    
    def _get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Get comprehensive file information"""
        try:
            stats = file_path.stat()
            file_hash = self._calculate_file_hash(file_path)
            
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_stem": file_path.stem,  
                "file_hash": file_hash,
                "file_size": stats.st_size,
                "modified_time": stats.st_mtime,
                "modified_datetime": datetime.fromtimestamp(stats.st_mtime).isoformat(),
                "extension": file_path.suffix.lower()
            }
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error getting file info: {e}")
            return {}
    
    def _check_duplicate_strategy(self, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Check how to handle a file based on duplicate strategy"""
        file_name = file_info["file_name"]
        file_stem = file_info["file_stem"]
        file_hash = file_info["file_hash"]
        
        # Check if exact same content exists
        exact_duplicate = False
        for processed_info in self.processed_files_log.values():
            if processed_info.get("file_hash") == file_hash:
                exact_duplicate = True
                break
        
        if exact_duplicate and self.duplicate_mode == DuplicateHandlingMode.SKIP:
            return {
                "action": "skip",
                "reason": "Exact duplicate found",
                "duplicate_info": processed_info
            }
        
        # Check for same filename with different content
        same_name_different_content = False
        previous_versions = []
        
        for key, processed_info in self.processed_files_log.items():
            if (processed_info.get("file_stem") == file_stem and 
                processed_info.get("file_hash") != file_hash):
                same_name_different_content = True
                previous_versions.append(processed_info)
        
        # Handle based on duplicate mode
        if same_name_different_content:
            if self.duplicate_mode == DuplicateHandlingMode.OVERWRITE:
                return {
                    "action": "overwrite",
                    "reason": "Overwrite mode enabled",
                    "previous_versions": previous_versions
                }
            elif self.duplicate_mode == DuplicateHandlingMode.VERSION:
                return {
                    "action": "version",
                    "reason": "Version mode enabled", 
                    "previous_versions": previous_versions
                }
            elif self.duplicate_mode == DuplicateHandlingMode.TIMESTAMP:
                latest_existing_time = max(
                    (v.get("modified_time", 0) for v in previous_versions),
                    default=0
                )
                if file_info["modified_time"] > latest_existing_time:
                    return {
                        "action": "overwrite",
                        "reason": "Newer timestamp detected",
                        "previous_versions": previous_versions
                    }
                else:
                    return {
                        "action": "skip",
                        "reason": "Older or same timestamp",
                        "previous_versions": previous_versions
                    }
        
        return {
            "action": "process",
            "reason": "New file",
            "previous_versions": []
        }
    
    def _remove_previous_versions(self, previous_versions: List[Dict[str, Any]]):
        """Remove previous versions from ChromaDB, BM25, and logs"""
        try:
            chunk_ids_to_remove = []
            bm25_indices_to_remove = []
            
            for version_info in previous_versions:
                document_id = version_info.get("document_id")
                if document_id:
                    # Remove from ChromaDB
                    try:
                        if self.collection is None:
                            logger.error("Collection not initialized")
                            continue
                        
                        results = self.collection.get(
                            where={"document_id": document_id},
                            include=["documents"]
                        )
                        if results and results.get("ids"):
                            chunk_ids_to_remove.extend(results["ids"])
                    except Exception as e:
                        logger.warning(f"Ã¢Å¡ Ã¯Â¸Â Could not find chunks for document {document_id}: {e}")
                    
                    # Remove from BM25 index
                    for i, metadata in enumerate(self.bm25_metadata):
                        if metadata.get("document_id") == document_id:
                            bm25_indices_to_remove.append(i)
            
            # Remove chunks from ChromaDB
            if chunk_ids_to_remove:
                try:
                    if self.collection is None:
                        logger.error("Collection not initialized")
                        return
                    self.collection.delete(ids=chunk_ids_to_remove)
                    logger.info(f"Ã°Å¸â€”â€˜Ã¯Â¸Â Removed {len(chunk_ids_to_remove)} chunks from ChromaDB")
                except Exception as e:
                    logger.error(f"Ã¢ÂÅ’ Error removing chunks: {e}")
            
            # Remove from BM25 index (reverse order to maintain indices)
            for i in sorted(bm25_indices_to_remove, reverse=True):
                if i < len(self.bm25_documents):
                    del self.bm25_documents[i]
                    del self.bm25_metadata[i]
            
            if bm25_indices_to_remove:
                # Rebuild BM25 index
                if self.bm25_documents:
                    tokenized_docs = [doc.lower().split() for doc in self.bm25_documents]
                    self.bm25_index = BM25Okapi(tokenized_docs)
            
            # Remove from processing log
            keys_to_remove = []
            for key, processed_info in self.processed_files_log.items():
                if any(processed_info.get("document_id") == v.get("document_id") 
                       for v in previous_versions):
                    keys_to_remove.append(key)
            
            for key in keys_to_remove:
                del self.processed_files_log[key]
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error removing previous versions: {e}")
    
    def process_all_documents(self) -> Dict[str, Any]:
        """Process all documents with Anthropic's Contextual Retrieval approach"""
        try:
            start_time = time.time()
            logger.info(f"Ã°Å¸â€Â Starting Anthropic-style contextual document processing...")
            
            # Get all supported files
            all_files = []
            for file_path in self.documents_dir.rglob("*"):
                if (file_path.is_file() and 
                    file_path.suffix.lower() in self.supported_extensions):
                    all_files.append(file_path)
            
            if not all_files:
                return {
                    "success": True,
                    "message": "No documents found",
                    "files_processed": 0,
                    "files_skipped": 0,
                    "files_overwritten": 0,
                    "processing_time": 0,
                    "contextual_enhancement": False,
                    "timestamp": datetime.now().isoformat()
                }
            
            # Process each file with contextual enhancement
            results = []
            processed_count = 0
            skipped_count = 0
            overwritten_count = 0
            error_count = 0
            contextual_enhancements = 0
            
            for file_path in all_files:
                try:
                    file_info = self._get_file_info(file_path)
                    if not file_info:
                        continue
                    
                    # Check duplicate handling strategy
                    strategy = self._check_duplicate_strategy(file_info)
                    
                    if strategy["action"] == "skip":
                        logger.info(f"Ã¢ÂÂ­Ã¯Â¸Â Skipping {file_path.name}: {strategy['reason']}")
                        skipped_count += 1
                        results.append({
                            "success": True,
                            "file": file_path.name,
                            "action": "skipped",
                            "reason": strategy["reason"]
                        })
                        continue
                    
                    elif strategy["action"] == "overwrite":
                        logger.info(f"Ã°Å¸â€â€ž Overwriting {file_path.name}: {strategy['reason']}")
                        self._remove_previous_versions(strategy["previous_versions"])
                        overwritten_count += 1
                    
                    elif strategy["action"] == "version":
                        logger.info(f"Ã°Å¸â€œÅ¡ Creating new version of {file_path.name}")
                    
                    # Process the file with contextual enhancement
                    result = self._process_single_file_contextual(file_path, file_info, strategy)
                    results.append(result)
                    
                    if result["success"]:
                        processed_count += 1
                        if result.get("contextual_enhanced"):
                            contextual_enhancements += 1
                    else:
                        error_count += 1
                        
                except Exception as e:
                    logger.error(f"Ã¢ÂÅ’ Error processing {file_path.name}: {e}")
                    error_count += 1
                    results.append({
                        "success": False,
                        "file": file_path.name,
                        "error": str(e)
                    })
            
            # Save logs and update statistics
            processing_time = time.time() - start_time
            self._update_processing_statistics(processed_count, skipped_count, overwritten_count, 
                                             error_count, processing_time, contextual_enhancements)
            self._save_processing_logs()
            
            logger.info(f"Ã¢Å“â€¦ Contextual processing complete - Processed: {processed_count}, "
                       f"Contextually Enhanced: {contextual_enhancements}, Skipped: {skipped_count}, "
                       f"Overwritten: {overwritten_count}, Errors: {error_count}")
            
            return {
                "success": True,
                "message": f"Anthropic-style contextual processing complete",
                "files_processed": processed_count,
                "files_skipped": skipped_count,
                "files_overwritten": overwritten_count,
                "files_with_errors": error_count,
                "contextual_enhancements": contextual_enhancements,
                "total_files_found": len(all_files),
                "processing_time": round(processing_time, 2),
                "retrieval_strategy": self.retrieval_strategy.value,
                "results": results,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error in contextual document processing: {e}")
            return {
                "success": False,
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }
    
    def _process_single_file_contextual(self, file_path: Path, file_info: Dict[str, Any], 
                                      strategy: Dict[str, Any]) -> Dict[str, Any]:
        """Process a single document file with Anthropic's contextual approach"""
        try:
            logger.info(f"Ã°Å¸â€œâ€ž Processing file with contextual enhancement: {file_path.name}")
            
            # Extract text based on file type
            extension = file_path.suffix.lower()
            if extension == '.pdf':
                text_content = self._extract_pdf_text(file_path)
            elif extension == '.docx':
                text_content = self._extract_docx_text(file_path)
            elif extension in ['.txt', '.md']:
                text_content = self._extract_text_file(file_path)
            else:
                return {
                    "success": False,
                    "file": file_path.name,
                    "error": f"Unsupported file type: {extension}"
                }
            
            # Process and chunk content
            processed_content = self._preprocess_text(text_content)
            # Extract title and status for metadata
            title, status = self._extract_title_and_status_from_text(processed_content, file_info["file_stem"])
            chunks = self._enhanced_project_aware_chunking(processed_content, file_path)
            
            # Generate document ID
            if strategy["action"] == "version":
                version_number = len(strategy["previous_versions"]) + 1
                doc_id = f"{self._generate_document_id(str(file_path))}_v{version_number}"
            else:
                doc_id = self._generate_document_id(str(file_path))
            
            # Prepare base metadata
            document_metadata = {
                "document_id": doc_id,
                "source_file": file_path.name,
                "title": title,  
                "status": status,
                "file_type": extension,
                "file_path": str(file_path.relative_to(self.documents_dir)),
                "file_hash": file_info["file_hash"],
                "file_stem": file_info["file_stem"],
                "ingestion_date": datetime.now().isoformat(),
                "modified_date": file_info["modified_datetime"],
                "total_chunks": len(chunks),
                "file_size": file_info["file_size"],
                "processing_version": "anthropic_contextual_v1.0",
                "duplicate_handling": self.duplicate_mode.value,
                "processing_action": strategy["action"],
                "chunking_strategy": self.chunking_strategy.value,
                "retrieval_strategy": self.retrieval_strategy.value,
                "embedding_model": self.embedding_model_name,
                "contextual_enhancement": self.claude_client is not None
            }
            
            # Add version info if applicable
            if strategy["action"] == "version":
                document_metadata["version_number"] = len(strategy["previous_versions"]) + 1
                document_metadata["is_latest_version"] = True
            
            # Apply contextual enhancement to chunks
            contextual_chunks = []
            contextual_enhanced = False
            total_chunks = len(chunks)
            
            logger.info(f"Ã°Å¸Â§  Generating contextual information for {total_chunks} chunks...")
            
            for i, chunk_data in enumerate(chunks):
                original_chunk = chunk_data["content"]

                if total_chunks > 5:  # Only show progress for larger documents
                    logger.info(f"Ã°Å¸â€œÂ Processing chunk {i+1}/{total_chunks}")
                
                # Generate contextual information (with rate limiting)
                contextual_chunk = self._generate_contextual_information(
                    original_chunk, processed_content, document_metadata
                )
                
                # Track if contextual enhancement was applied
                if len(contextual_chunk) > len(original_chunk):
                    contextual_enhanced = True
                
                contextual_chunks.append({
                    **chunk_data,
                    "original_content": original_chunk,
                    "contextual_content": contextual_chunk,
                    "context_added": len(contextual_chunk) - len(original_chunk)
                })
            
            # Store chunks in vector database and BM25 index
            self._store_chunks_contextual(contextual_chunks, document_metadata, doc_id)
            
            # Store chunks with enhanced metadata
            self._store_chunks_with_project_metadata(chunks, document_metadata, doc_id)
            
            # Record processing success with project count
            project_count = len([c for c in chunks if c.get("project_metadata", {}).get("project_name")])
           
            # Record processing success
            processing_record = {
                **file_info,
                "processed_date": datetime.now().isoformat(),
                "document_id": doc_id,
                "chunks_created": len(chunks),
                "projects_identified": project_count,
                "contextual_enhanced": True,
                "chunking_method": "project_aware",
                "status": "success"
            }
            
            # Add version info
            if strategy["action"] == "version":
                processing_record["version_number"] = document_metadata["version_number"]
                processing_record["is_latest_version"] = True
            
            # Update file versions tracking
            self._update_file_versions(file_info["file_stem"], doc_id, file_info["file_hash"], processing_record)
            
            self.processed_files_log[file_path.name] = processing_record
            
            logger.info(f"Ã¢Å“â€¦ Successfully processed with {project_count} projects identified: {file_path.name}")
            
            return {
                "success": True,
                "file": file_path.name,
                "document_id": doc_id,
                "chunks_created": len(chunks),
                "projects_identified": project_count,
                "contextual_enhanced": True,
                "action": strategy["action"]
            }
                
        except Exception as e:
            logger.error(f"Error processing file {file_path.name}: {e}")
            return {
                "success": False,
                "file": file_path.name,
                "error": str(e)
            }
    
    def _store_chunks_with_project_metadata(self, chunks: List[Dict[str, Any]], 
                                       document_metadata: Dict[str, Any], doc_id: str):
        """Store chunks with enhanced project metadata"""
        
        if not self.embedding_model or not self.collection:
            raise RuntimeError("Embedding model or collection not initialized")
        
        chunk_ids = []
        chunk_metadatas = []
        chunk_documents = []
        chunk_embeddings = []
        
        for i, chunk_data in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            
            # Use contextual content for embedding and storage
            contextual_content = chunk_data.get("contextual_content", chunk_data["content"])
            original_content = chunk_data.get("original_content", chunk_data["content"])
            
            # Generate embedding from contextual content
            embedding = self.embedding_model.encode(contextual_content).tolist()
            
            # Enhanced metadata with project information
            chunk_metadata = {
                **document_metadata,
                "chunk_index": i,
                "chunk_id": chunk_id,
                "chunk_type": chunk_data.get("type", "content"),
                "chunk_size": chunk_data.get("size", len(original_content)),
                "optimized_for_context": chunk_data.get("optimized_for_context", False)
            }

            # Add project-specific metadata if available
            if "project_metadata" in chunk_data:
                project_meta = chunk_data["project_metadata"]
                chunk_metadata.update({
                    "project_name": project_meta.get("project_name"),
                    "project_status": project_meta.get("status"),
                    "product_owner": project_meta.get("product_owner"),
                    "is_project_summary": project_meta.get("is_summary", False),
                    "total_projects_in_doc": project_meta.get("total_projects"),
                    "has_complete_project_info": project_meta.get("has_complete_info", False)
                })
            
            chunk_ids.append(chunk_id)
            chunk_metadatas.append(chunk_metadata)
            chunk_documents.append(contextual_content)
            chunk_embeddings.append(embedding)
            
            # Add to BM25 index for hybrid search
            self.bm25_documents.append(contextual_content)
            self.bm25_metadata.append(chunk_metadata)
        
        # Add to ChromaDB collection
        self.collection.add(
            documents=chunk_documents,
            metadatas=chunk_metadatas,
            ids=chunk_ids,
            embeddings=chunk_embeddings
        )
        
        # Rebuild BM25 index if we have documents
        if self.bm25_documents:
            tokenized_docs = [doc.lower().split() for doc in self.bm25_documents]
            self.bm25_index = BM25Okapi(tokenized_docs)
    def query_documents_with_project_awareness(self, query: str, top_k: int = 8) -> Dict[str, Any]:
        """Enhanced query with project-specific awareness"""
        
        try:
            # Determine if this is a project-specific query
            query_analysis = self._analyze_project_query(query)
            
            if query_analysis["is_project_specific"]:
                # Use project-specific retrieval
                return self._query_specific_project(
                    query, 
                    query_analysis["project_name"], 
                    top_k
                )
            elif query_analysis["is_status_query"]:
                # Use status-based retrieval
                return self._query_by_project_status(
                    query,
                    query_analysis["target_status"],
                    top_k
                )
            else:
                # Use standard contextual hybrid search
                return self._query_contextual_hybrid(query, top_k)
                
        except Exception as e:
            logger.error(f"Project-aware query failed: {e}")
            return self._query_contextual_hybrid(query, top_k)

    def _analyze_project_query(self, query: str) -> Dict[str, Any]:
        """Analyze if query is project-specific or status-based"""
        
        query_lower = query.lower()
        
        # Check for specific project names
        project_names = [
            "refinancing marketing filter", "data masking", "fraud detection",
            "loan eligibility", "portfolio analyzer", "compliance reporting",
            "predictive maintenance", "financial advice", "customer onboarding",
            "payment gateway"
        ]
        
        for project in project_names:
            if project in query_lower:
                return {
                    "is_project_specific": True,
                    "project_name": project,
                    "is_status_query": False,
                    "target_status": None
                }
        
        # Check for status-based queries
        status_patterns = {
            "incomplete": ["in progress", "halted"],
            "unfinished": ["in progress", "halted"],
            "current": ["in progress"],
            "ongoing": ["in progress"],
            "halted": ["halted"],
            "stopped": ["halted"],
            "completed": ["completed"],
            "finished": ["completed"]
        }
        
        for query_term, target_statuses in status_patterns.items():
            if query_term in query_lower:
                return {
                    "is_project_specific": False,
                    "project_name": None,
                    "is_status_query": True,
                    "target_status": target_statuses
                }
        
        return {
            "is_project_specific": False,
            "project_name": None,
            "is_status_query": False,
            "target_status": None
        }

    def _query_specific_project(self, query: str, project_name: str, top_k: int) -> Dict[str, Any]:
        """Query for a specific project"""
        
        if not self.collection:
            return {"success": False, "error": "Collection not initialized"}
        
        try:
            # First, try to find project-specific chunks using correct ChromaDB where syntax
            results = self.collection.query(
                query_texts=[query],
                n_results=top_k * 2,
                where={"project_name": {"$eq": project_name}},  # Fixed: correct ChromaDB syntax
                include=["documents", "metadatas", "distances"]
            )
            
            # If no project-specific results, fall back to text search
            if not results["documents"] or not results["documents"][0]:
                results = self.collection.query(
                    query_texts=[f"{project_name} {query}"],
                    n_results=top_k,
                    include=["documents", "metadatas", "distances"]
                )
            
            # Convert ChromaDB QueryResult to Dict format for _format_query_results
            results_dict = {
                "documents": results.get("documents", []),
                "metadatas": results.get("metadatas", []),
                "distances": results.get("distances", []),
                "ids": results.get("ids", [])
            }
            
            return self._format_query_results(results_dict, query, "project_specific")
            
        except Exception as e:
            logger.error(f"Error in project-specific query: {e}")
            return {"success": False, "error": str(e)}

    def _query_by_project_status(self, query: str, target_statuses: List[str], top_k: int) -> Dict[str, Any]:
        """Query projects by status"""
        
        if not self.collection:
            return {"success": False, "error": "Collection not initialized"}
        
        try:
            # Try to get summary chunk first
            summary_results = self.collection.query(
                query_texts=[query],
                n_results=5,
                where={"is_project_summary": {"$eq": True}},  # Fixed: correct syntax
                include=["documents", "metadatas", "distances"]
            )
            
            # Get individual project chunks with matching status
            status_where = {
                "$or": [
                    {"project_status": {"$eq": status}}
                    for status in target_statuses
                ]
            }

            project_results = self.collection.query(
                query_texts=[query],
                n_results=top_k,
                where=status_where,  # type: ignore
                include=["documents", "metadatas", "distances"]
            )
            
            # Safely combine results
            combined_docs = []
            combined_metadatas = []
            combined_distances = []
            combined_ids = []
            
            # Add summary results
            if summary_results["documents"] and summary_results["documents"][0]:
                combined_docs.extend(summary_results["documents"][0])
                assert summary_results["metadatas"] is not None
                combined_metadatas.extend(summary_results["metadatas"][0])
                assert summary_results["distances"] is not None
                combined_distances.extend(summary_results["distances"][0])
                assert summary_results["ids"] is not None
                combined_ids.extend(summary_results["ids"][0])
            
            # Add project results
            if project_results["documents"] and project_results["documents"][0]:
                combined_docs.extend(project_results["documents"][0])
                assert project_results["metadatas"] is not None
                combined_metadatas.extend(project_results["metadatas"][0])
                assert project_results["distances"] is not None
                combined_distances.extend(project_results["distances"][0])
                assert project_results["ids"] is not None
                combined_ids.extend(project_results["ids"][0])
            
            # Create properly formatted results
            combined_results = {
                "documents": [combined_docs],
                "metadatas": [combined_metadatas],
                "distances": [combined_distances],
                "ids": [combined_ids]
            }
            
            return self._format_query_results(combined_results, query, "status_filtered")
            
        except Exception as e:
            logger.error(f"Error in status-based query: {e}")
            return {"success": False, "error": str(e)}

    def _extract_pdf_text(self, pdf_path: Path) -> str:
        """Extract text content from PDF file"""
        try:
            text_content = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n\n"
            
            if not text_content.strip():
                raise ValueError("No text content extracted from PDF")
            
            return text_content
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error extracting PDF text: {e}")
            raise
    
    def _extract_docx_text(self, docx_path: Path) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = Document(str(docx_path))
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content += f"{text}\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content += f"{' | '.join(row_text)}\n"
            
            if not text_content.strip():
                raise ValueError("No text content extracted from DOCX")
            
            return text_content
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error extracting DOCX text: {e}")
            raise
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text content from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            if not text_content.strip():
                raise ValueError("No text content in file")
            
            return text_content
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error extracting text file: {e}")
            raise
    
    def _extract_title_and_status_from_text(self, text: str, file_stem: Optional[str] = None) -> tuple[str, str]:
        """
        Heuristic extraction of a project title and Status from a document.
        Returns (title, status). Falls back to file_stem / defaults.
        """
        import re
        title = None
        status = "unknown"

        # quick normalization
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

        # 1) Numbered heading e.g. "4. Loan Eligibility Predictor"
        for ln in lines[:60]:
            m = re.match(r'^\s*\d+\.\s*(.+)', ln)
            if m:
                title = m.group(1).strip()
                break

        # 2) explicit title lines: e.g., "Loan Eligibility Predictor" as a short line near top
        if not title:
            for ln in lines[:12]:
                if re.search(r'^(Status|Description|Business Objective)\b', ln, re.I):
                    continue
                if 2 <= len(ln.split()) <= 8 and not ln.endswith(':'):
                    title = ln.strip()
                    break

        # 3) extract Status (searched through first ~60 lines)
        for ln in lines[:80]:
            m = re.search(r'\bStatus\s*[:\-\Ã¢â‚¬â€œ]\s*(.+)', ln, re.I)
            if m:
                status = m.group(1).strip()
                break
            m2 = re.search(r'\bStatus\b.*\b(In Progress|Completed|Halted|On Hold|In Review|Draft|Planned)\b', ln, re.I)
            if m2:
                status = m2.group(1).strip()
                break

        # final fallback
        title = (title or file_stem or "Untitled").strip()
        status = (status or "unknown").strip()
        return title, status


    def _preprocess_text(self, text: str) -> str:
        """Preprocess text content"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Clean up common artifacts
        text = text.replace('\x00', '')
        text = re.sub(r'[\r\f\v]', ' ', text)
        
        return text.strip()
    
    def _chunk_text(self, text: str, file_path: Path) -> List[Dict[str, Any]]:
        """Chunk text using the specified strategy with Anthropic's recommended settings"""
        if self.chunking_strategy == ChunkingStrategy.SEMANTIC:
            return self._chunk_semantic(text)
        elif self.chunking_strategy == ChunkingStrategy.SENTENCE:
            return self._chunk_by_sentences(text)
        elif self.chunking_strategy == ChunkingStrategy.PARAGRAPH:
            return self._chunk_by_paragraphs(text)
        else:  # FIXED_SIZE
            return self._chunk_fixed_size(text)
    
    def _chunk_semantic(self, text: str) -> List[Dict[str, Any]]:
        """Semantic-aware chunking optimized for Anthropic's approach"""
        chunks = []
        
        # Split by double newlines (paragraphs) first
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_size = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            paragraph_size = len(paragraph)
            
            # If paragraph alone exceeds max size, split it further
            if paragraph_size > self.max_chunk_size:
                # Save current chunk if exists
                if current_chunk:
                    chunks.append({
                        "content": current_chunk.strip(),
                        "type": "semantic",
                        "size": current_size,
                        "optimized_for_context": True
                    })
                    current_chunk = ""
                    current_size = 0
                
                # Split large paragraph by sentences
                sentences = re.split(r'[.!?]+', paragraph)
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                        
                    sentence_with_punct = sentence + ". "
                    if current_size + len(sentence_with_punct) > self.chunk_size and current_chunk:
                        chunks.append({
                            "content": current_chunk.strip(),
                            "type": "semantic",
                            "size": current_size,
                            "optimized_for_context": True
                        })
                        current_chunk = sentence_with_punct
                        current_size = len(sentence_with_punct)
                    else:
                        current_chunk += sentence_with_punct
                        current_size += len(sentence_with_punct)
            else:
                # Check if adding this paragraph exceeds chunk size
                if current_size + paragraph_size > self.chunk_size and current_chunk:
                    chunks.append({
                        "content": current_chunk.strip(),
                        "type": "semantic",
                        "size": current_size,
                        "optimized_for_context": True
                    })
                    current_chunk = paragraph + "\n\n"
                    current_size = paragraph_size + 2
                else:
                    current_chunk += paragraph + "\n\n"
                    current_size += paragraph_size + 2
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                "content": current_chunk.strip(),
                "type": "semantic",
                "size": current_size,
                "optimized_for_context": True
            })
        
        return chunks
    
    def _chunk_by_sentences(self, text: str) -> List[Dict[str, Any]]:
        """Chunk text by sentences"""
        chunks = []
        sentences = re.split(r'[.!?]+', text)
        
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_with_punct = sentence + ". "
            sentence_size = len(sentence_with_punct)
            
            if current_size + sentence_size > self.chunk_size and current_chunk:
                chunks.append({
                    "content": current_chunk.strip(),
                    "type": "sentence",
                    "size": current_size,
                    "optimized_for_context": True
                })
                current_chunk = sentence_with_punct
                current_size = sentence_size
            else:
                current_chunk += sentence_with_punct
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                "content": current_chunk.strip(),
                "type": "sentence",
                "size": current_size,
                "optimized_for_context": True
            })
        
        return chunks
    
    def _chunk_by_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """Chunk text by paragraphs"""
        chunks = []
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_size = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            paragraph_with_break = paragraph + "\n\n"
            paragraph_size = len(paragraph_with_break)
            
            if paragraph_size > self.max_chunk_size:
                # Split large paragraph further
                if current_chunk:
                    chunks.append({
                        "content": current_chunk.strip(),
                        "type": "paragraph",
                        "size": current_size,
                        "optimized_for_context": True
                    })
                    current_chunk = ""
                    current_size = 0
                
                # Use sentence chunking for large paragraphs
                para_chunks = self._chunk_by_sentences(paragraph)
                chunks.extend(para_chunks)
            else:
                if current_size + paragraph_size > self.chunk_size and current_chunk:
                    chunks.append({
                        "content": current_chunk.strip(),
                        "type": "paragraph",
                        "size": current_size,
                        "optimized_for_context": True
                    })
                    current_chunk = paragraph_with_break
                    current_size = paragraph_size
                else:
                    current_chunk += paragraph_with_break
                    current_size += paragraph_size
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                "content": current_chunk.strip(),
                "type": "paragraph",
                "size": current_size,
                "optimized_for_context": True
            })
        
        return chunks
    
    def _chunk_fixed_size(self, text: str) -> List[Dict[str, Any]]:
        """Traditional fixed-size chunking with overlap (Anthropic optimized)"""
        chunks = []
        words = text.split()
        
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            
            if current_length + word_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunk_content = " ".join(current_chunk)
                chunks.append({
                    "content": chunk_content,
                    "type": "fixed_size",
                    "size": len(chunk_content),
                    "optimized_for_context": True
                })
                
                # Start new chunk with overlap
                overlap_words = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                current_chunk = overlap_words + [word]
                current_length = sum(len(w) + 1 for w in current_chunk)
            else:
                current_chunk.append(word)
                current_length += word_length
        
        # Add final chunk
        if current_chunk:
            chunk_content = " ".join(current_chunk)
            chunks.append({
                "content": chunk_content,
                "type": "fixed_size",
                "size": len(chunk_content),
                "optimized_for_context": True
            })
        
        return chunks
    
    def _store_chunks_contextual(self, chunks: List[Dict[str, Any]], 
                               document_metadata: Dict[str, Any], doc_id: str):
        """Store chunks in both ChromaDB and BM25 index with contextual enhancement"""
        if not self.embedding_model or not self.collection:
            raise RuntimeError("Embedding model or collection not initialized")
        
        chunk_ids = []
        chunk_metadatas = []
        chunk_documents = []
        chunk_embeddings = []
        
        for i, chunk_data in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            
            # Use contextual content for embedding and storage
            contextual_content = chunk_data["contextual_content"]
            original_content = chunk_data["original_content"]
            
            # Generate embedding from contextual content
            embedding = self.embedding_model.encode(contextual_content).tolist()
            
            chunk_metadata = {
                **document_metadata,
                "chunk_index": i,
                "chunk_id": chunk_id,
                "chunk_type": chunk_data.get("type", "content"),
                "chunk_size": chunk_data.get("size", len(original_content)),
                "context_added": chunk_data.get("context_added", 0),
                "has_context": chunk_data.get("context_added", 0) > 0,
                "optimized_for_context": chunk_data.get("optimized_for_context", False)
            }
            
            chunk_ids.append(chunk_id)
            chunk_metadatas.append(chunk_metadata)
            chunk_documents.append(contextual_content)  # Store contextual version
            chunk_embeddings.append(embedding)
            
            # Add to BM25 index for hybrid search
            self.bm25_documents.append(contextual_content)
            self.bm25_metadata.append(chunk_metadata)
        
        # Add to ChromaDB collection
        self.collection.add(
            documents=chunk_documents,
            metadatas=chunk_metadatas,
            ids=chunk_ids,
            embeddings=chunk_embeddings
        )
        
        # Rebuild BM25 index if we have documents
        if self.bm25_documents:
            tokenized_docs = [doc.lower().split() for doc in self.bm25_documents]
            self.bm25_index = BM25Okapi(tokenized_docs)
    
    def query_documents(self, query: str, top_k: int = 6, 
                       retrieval_strategy: Optional[RetrievalStrategy] = None) -> Dict[str, Any]:
        """
        Query documents using Anthropic's Contextual Retrieval approach
        
        Args:
            query: Search query
            top_k: Number of results to return
            retrieval_strategy: Override default retrieval strategy
            
        Returns:
            Query results with contextual information
        """
        try:
            if not self.embedding_model or not self.collection:
                return {
                    "success": False,
                    "error": "Embedding model or collection not initialized"
                }
            
            strategy = retrieval_strategy or self.retrieval_strategy
            
            if strategy == RetrievalStrategy.VECTOR_ONLY:
                return self._query_vector_only(query, top_k)
            elif strategy == RetrievalStrategy.BM25_ONLY:
                return self._query_bm25_only(query, top_k)
            elif strategy == RetrievalStrategy.HYBRID:
                return self._query_hybrid(query, top_k)
            elif strategy == RetrievalStrategy.CONTEXTUAL_HYBRID:
                return self._query_contextual_hybrid(query, top_k)
            else:
                return self._query_vector_only(query, top_k)
                
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error querying documents: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _format_results_fast(self, results: Dict[str, Any], query: str) -> Dict[str, Any]:
        """CORRECTED: Add relevance filtering to remove irrelevant chunks"""
        try:
            if not results or not results.get("documents"):
                return {"success": True, "query": query, "documents": [], "message": "No documents found"}
            
            documents = results["documents"][0] if results["documents"] else []
            metadatas = results["metadatas"][0] if results["metadatas"] else []
            distances = results["distances"][0] if results["distances"] else []
            
            formatted_results = []
            min_length = min(len(documents), len(metadatas), len(distances))
            
            for i in range(min_length):
                distance = distances[i] if i < len(distances) else 1.0
                similarity_score = max(0, 1 - min(distance, 1.0))
                
                # Use the existing _extract_actual_content method instead
                content = self._extract_actual_content(documents[i])
                metadata = metadatas[i] if i < len(metadatas) else {}
                
                # NEW: Add relevance check
                if self._is_content_relevant(content, query):
                    formatted_results.append({
                        "content": content,
                        "metadata": metadata,
                        "similarity_score": similarity_score,
                        "source_file": metadata.get("source_file", "Unknown"),
                        "title": metadata.get("title"),
                        "status": metadata.get("status")
                    })
            
            return {
                "success": True,
                "query": query,
                "documents": formatted_results,
                "total_results": len(formatted_results)
            }
            
        except Exception as e:
            logger.error(f"Error formatting results: {e}")
            return {"success": False, "error": str(e)}

    def _is_content_relevant(self, content: str, query: str) -> bool:
        """NEW METHOD: Check if content is relevant to the query"""
        query_lower = query.lower()
        content_lower = content.lower()
        
        # For follow-up questions, check for specific project mentions
        if any(word in query_lower for word in ["refinancing", "marketing", "filter"]):
            # Must contain these specific terms
            return "refinancing marketing filter" in content_lower
        
        if any(word in query_lower for word in ["data", "masking"]):
            return "data masking tool" in content_lower and "data masking tool" in content_lower[:200]
        
        if any(word in query_lower for word in ["fraud", "detection"]):
            return "fraud detection" in content_lower
        
        # For general queries, use basic relevance
        query_words = set(query_lower.split())
        content_words = set(content_lower.split())
        overlap = len(query_words.intersection(content_words))
        
        # At least 30% word overlap or high similarity
        return overlap >= len(query_words) * 0.3

    def _enhance_query_for_retrieval(self, query: str) -> Optional[str]:
        """AI-powered semantic query enhancement without hardcoded patterns"""
        
        if not self.claude_client:
            return None
        
        prompt = f"""Enhance this query for better document retrieval:

    Original Query: "{query}"

    Analyze the user's intent and add relevant search terms that would appear in business documents.
    Consider synonyms, related concepts, and alternative phrasings that might be used in formal documentation.

    Return only the enhanced query, or "NONE" if the original is sufficient.

    Enhanced Query:"""

        try:
            def make_claude_call():
                if not self.claude_client:
                    return None
                return self.claude_client.messages.create(
                    model="claude-3-5-haiku-20241022",
                    max_tokens=100,
                    temperature=0.1,
                    messages=[{"role": "user", "content": prompt}]
                )
            
            response = self._rate_limited_api_call(make_claude_call)
            
            if response:
                enhanced = "".join([
                    block.text for block in response.content 
                    if hasattr(block, 'text')
                ]).strip()
                
                return None if enhanced.upper() == "NONE" else enhanced
                
        except Exception as e:
            logger.warning(f"Semantic query enhancement failed: {e}")
        
        return None

    def _query_vector_only(self, query: str, top_k: int) -> Dict[str, Any]:
        """Query using vector similarity only"""
        try:
            # Generate query embedding
            if self.embedding_model is None:
                return {
                    "success": False, 
                    "error": "Embedding model not initialized. Please check your sentence-transformers installation."
                }
            
            # FIX: Convert numpy array to Python list for ChromaDB
            query_embedding_raw = self.embedding_model.encode([query], show_progress_bar=False)
            
            # Ensure it's a Python list, not numpy array
            if hasattr(query_embedding_raw, 'tolist'):
                query_embedding = query_embedding_raw[0].tolist()  # Convert to list
            else:
                query_embedding = list(query_embedding_raw[0])  # Fallback conversion
            
            # Query ChromaDB
            if self.collection is None:
                return {
                    "success": False, 
                    "error": "ChromaDB collection not initialized"
                }
            results = self.collection.query(
                query_embeddings=[query_embedding],  # Now properly formatted as list
                n_results=top_k,
                include=["documents", "metadatas", "distances"]
            )

            results_dict = {
                "documents": results.get("documents", []),
                "metadatas": results.get("metadatas", []),
                "distances": results.get("distances", []),
                "ids": results.get("ids", [])
            }
            
            # Format the results
            formatted_results = self._format_query_results(results_dict, query, "vector_only")
            
            # Add direct answer extraction using AI
            if formatted_results.get("success") and formatted_results.get("documents"):
                direct_answer = self.extract_direct_answer(query, formatted_results["documents"])
                formatted_results["direct_answer"] = direct_answer
            
            return formatted_results
        
        except Exception as e:
            logger.error(f"Error in vector query: {e}")
            return {"success": False, "error": str(e)}
    
    def _enhance_query_with_ai(self, query: str) -> str:
        """Use AI to enhance query for better retrieval"""
        if not self.claude_client:
            return query
        
        prompt = f"""Enhance this search query for better document retrieval in a corporate banking and fintech environment.

    Original Query: "{query}"

    Add relevant business terms, synonyms, and alternative phrasings that might appear in:
    - Financial documents (policies, procedures, agreements)
    - Technical documentation (APIs, systems, configurations)  
    - Administrative records (contacts, roles, responsibilities)
    - Contract documents (terms, conditions, obligations)

    Provide an enhanced query with 5-10 additional relevant terms.
    Respond with just the enhanced query, nothing else."""

        try:
            def make_claude_call():
                if self.claude_client is None:
                    raise ValueError("Claude client not initialized")
                return self.claude_client.messages.create(
                    model="claude-3-5-haiku-20241022",
                    max_tokens=100,
                    temperature=0.1,
                    messages=[{"role": "user", "content": prompt}]
                )
            
            response = self._rate_limited_api_call(make_claude_call)
            
            if response is not None:
                # Handle the response properly
                enhanced_query = ""
                for block in response.content:
                    if hasattr(block, 'text'):
                        enhanced_query += block.text
                    elif isinstance(block, str):
                        enhanced_query += block
                
                if enhanced_query.strip():
                    return enhanced_query.strip()
        
        except Exception as e:
            logger.warning(f"Ã¢Å¡ Ã¯Â¸Â Query enhancement failed: {e}")
        
        return query

    
    def extract_direct_answer(self, query: str, search_results: List[Dict[str, Any]]) -> str:
        """AI-powered direct answer extraction with semantic understanding"""
        if not self.claude_client or not search_results:
            return "Information not found in the documents."
        
        # Combine top results with better context
        context = "\n\n---\n\n".join([
            f"Source: {r['source_file']}\nTitle: {r.get('title', 'N/A')}\nStatus: {r.get('status', 'N/A')}\nContent: {r['content'][:600]}"
            for r in search_results[:3]
        ])
        
        prompt = f"""Based on these document excerpts, provide a direct, accurate answer to the user's question.

    Question: "{query}"

    Document Context:
    {context}

    Instructions:
    1. Provide a specific, factual answer based ONLY on the provided documents
    2. If the answer involves structured information (names, dates, amounts, statuses), quote them exactly
    3. If no specific answer exists, say "The provided documents don't contain specific information about this question"
    4. For multi-part questions, address each part if information is available
    5. Keep the response concise and directly relevant to the question

    Answer:"""

        try:
            def make_claude_call():
                if not self.claude_client:
                    return ""
                return self.claude_client.messages.create(
                    model="claude-3-5-haiku-20241022",
                    max_tokens=300,
                    temperature=0.1,
                    messages=[{"role": "user", "content": prompt}]
                )
            
            response = self._rate_limited_api_call(make_claude_call)
            
            if response:
                answer = "".join([
                    block.text for block in response.content 
                    if hasattr(block, 'text')
                ]).strip()
                
                return answer if answer else "Information not available in the provided documents."
            
        except Exception as e:
            logger.warning(f"Direct answer extraction failed: {e}")
        
        return "Unable to generate answer from the provided documents."
    
    def _filter_and_extract_relevant_content(self, documents: List[Dict], query: str) -> List[Dict]:
        """Enhanced content filtering and extraction based on query intent"""
        
        # Add null safety check first
        if not self.claude_client or not documents:
            return documents
        
        # Analyze query to determine extraction intent
        extraction_prompt = f"""Analyze this query to determine what specific information should be extracted:

    Query: "{query}"

    Determine the extraction intent:
    1. PROJECT_OVERVIEW - user wants summary/description of specific project
    2. PROJECT_LIST_BY_STATUS - user wants list of projects with specific status  
    3. PROJECT_FEATURES - user wants detailed features/capabilities
    4. PROJECT_REASONS - user wants reasons for status (halted, etc.)

    Respond with just the intent type and target (if specific project mentioned):
    Format: "INTENT_TYPE|target_project_name" or just "INTENT_TYPE" """

        try:
            # Fixed: Use the _rate_limited_api_call method correctly with proper Claude client check
            def make_claude_call():
                if self.claude_client is None:
                    raise ValueError("Claude client not initialized")
                return self.claude_client.messages.create(
                    model="claude-3-5-haiku-20241022",
                    max_tokens=100,
                    temperature=0.1,
                    messages=[{"role": "user", "content": extraction_prompt}]
                )
            
            response = self._rate_limited_api_call(make_claude_call, "content_filtering")
            
            if not response:
                return documents
                
            intent_text = "".join([
                block.text for block in response.content 
                if hasattr(block, 'text')
            ]).strip()
            
            # Filter documents based on intent
            return self._apply_content_filtering(documents, intent_text, query)
            
        except Exception as e:
            logger.warning(f"Content filtering failed: {e}")
            return documents

    def _apply_content_filtering(self, documents: List[Dict], intent_text: str, query: str) -> List[Dict]:
        """Apply intelligent content filtering based on intent"""
        
        filtered_docs = []
        
        for doc in documents:
            content = doc.get("content", "")
            
            if "|" in intent_text:
                intent, target = intent_text.split("|", 1)
                target = target.strip().lower()
                
                # For specific project queries, only include content about that project
                if target in content.lower():
                    # Extract only the relevant section
                    filtered_content = self._extract_project_section(content, target)
                    if filtered_content:
                        doc_copy = doc.copy()
                        doc_copy["content"] = filtered_content
                        doc_copy["filtered"] = True
                        filtered_docs.append(doc_copy)
            else:
                # For general queries, apply intent-based filtering
                intent = intent_text.strip()
                
                if intent == "PROJECT_LIST_BY_STATUS":
                    # Include content that has clear status indicators
                    if any(status in content.lower() for status in ["status:", "in progress", "completed", "halted"]):
                        filtered_docs.append(doc)
                else:
                    filtered_docs.append(doc)
        
        return filtered_docs if filtered_docs else documents

    def _enhanced_project_aware_chunking(self, text: str, file_path: Path) -> List[Dict[str, Any]]:
        """Simplified project-aware chunking focused on the core issue"""
        
        # Simple project detection using clear markers
        lines = text.split('\n')
        project_boundaries = []
        
        # Find clear project boundaries
        for i, line in enumerate(lines):
            line_clean = line.strip()
            # Look for numbered projects or clear status markers
            if (re.match(r'^\d+\.', line_clean) or 
                line_clean.startswith('**') and ('Status:' in ''.join(lines[i:i+5]))):
                project_boundaries.append(i)
        
        chunks = []
        
        if len(project_boundaries) < 2:
            # If no clear projects found, use standard chunking
            return self._chunk_text(text, file_path)
        
        # Create one chunk per project section
        for i in range(len(project_boundaries)):
            start = project_boundaries[i]
            end = project_boundaries[i + 1] if i + 1 < len(project_boundaries) else len(lines)
            
            project_lines = lines[start:end]
            project_content = '\n'.join(project_lines)
            
            # Extract basic project info
            project_name = "Unknown Project"
            project_status = "unknown"
            
            for line in project_lines[:10]:  # Check first 10 lines only
                # Extract project name from first line
                if i == 0 or line.strip().startswith(('**', f'{i}.')):
                    clean_line = re.sub(r'^\d+\.\s*|\*\*', '', line.strip())
                    if clean_line and len(clean_line) < 100:
                        project_name = clean_line.replace('**', '').strip()
                
                # Extract status
                status_match = re.search(r'Status:\s*(In Progress|Completed|Halted)', line)
                if status_match:
                    project_status = status_match.group(1)
                    break
            
            # Only create separate chunks if content is substantial
            if len(project_content.strip()) > 50:
                chunks.append({
                    "content": project_content,
                    "type": "project_section",
                    "size": len(project_content),
                    "optimized_for_context": True,
                    "project_metadata": {
                        "project_name": project_name,
                        "status": project_status,
                        "has_complete_info": "Status:" in project_content
                    }
                })
        
        return chunks if chunks else self._chunk_text(text, file_path)

    def _identify_projects_in_document(self, text: str) -> List[Dict[str, Any]]:
        """Identify individual projects and extract complete information"""
        
        projects = []
        lines = text.split('\n')
        current_project = None
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Detect project start (numbered projects or clear project headers)
            project_match = re.match(r'^(\d+)\.\s*(.+)', line_stripped)
            if project_match:
                # Save previous project if exists
                if current_project:
                    projects.append(current_project)
                
                # Start new project
                project_number = project_match.group(1)
                project_name = project_match.group(2)
                
                current_project = {
                    "number": project_number,
                    "name": project_name,
                    "start_line": i,
                    "content_lines": [line],
                    "status": "unknown",
                    "product_owner": None,
                    "is_complete": False
                }
            
            # Detect project by status patterns (for projects without numbers)
            elif re.search(r'Status:\s*(In Progress|Completed|Halted)', line_stripped):
                m = re.search(r'Status:\s*(In Progress|Completed|Halted)', line_stripped)
                if m:
                    status = m.group(1)
                else:
                    status = "unknown"

                if not current_project:
                    # Look backward for project name
                    project_name = self._find_project_name_backward(lines, i)
                    current_project = {
                        "number": None,
                        "name": project_name,
                        "start_line": max(0, i-2),
                        "content_lines": lines[max(0, i-2):i+1],
                        "status": status,
                        "product_owner": None,
                        "is_complete": False
                    }
                else:
                    current_project["status"] = status
                    current_project["content_lines"].append(line)
            
            # Continue collecting content for current project
            elif current_project:
                current_project["content_lines"].append(line)
                
                # Extract product owner if found
                if "Product Owner:" in line:
                    owner_match = re.search(r'Product Owner:\s*(.+)', line)
                    if owner_match:
                        current_project["product_owner"] = owner_match.group(1).strip()
                
                # Check if we've reached next project or end
                if (i < len(lines) - 1 and 
                    (re.match(r'^\d+\.', lines[i+1].strip()) or 
                    lines[i+1].strip().startswith('**') and 'Status:' in ''.join(lines[i+1:i+5]))):
                    # End current project
                    projects.append(current_project)
                    current_project = None
        
        # Add final project if exists
        if current_project:
            projects.append(current_project)
        
        # Process projects to create complete content
        processed_projects = []
        for project in projects:
            full_content = '\n'.join(project["content_lines"])
            
            # Determine if project info is complete
            has_description = "Description:" in full_content
            has_status = project["status"] != "unknown"
            has_owner = project["product_owner"] is not None
            
            processed_project = {
                **project,
                "full_content": full_content,
                "is_complete": has_description and has_status,
                "content_quality": "high" if (has_description and has_status and has_owner) else "medium"
            }
            processed_projects.append(processed_project)
        
        return processed_projects

    def _find_project_name_backward(self, lines: List[str], current_index: int) -> str:
        """Find project name by looking backward from status line"""
        
        for i in range(current_index - 1, max(0, current_index - 5), -1):
            line = lines[i].strip()
            
            # Look for project names in headers or titles
            if line.startswith('**') and line.endswith('**'):
                return line.replace('**', '').strip()
            elif re.match(r'^\d+\.\s*(.+)', line):
                m = re.match(r'^\d+\.\s*(.+)', line)
                if m:
                    return m.group(1).strip()
            elif len(line) > 10 and len(line) < 100 and not line.endswith(':'):
                return line
        
        return "Unknown Project"

    def _create_project_summary(self, projects: List[Dict]) -> str:
        """Create comprehensive project summary for better retrieval"""
        
        summary_parts = ["PROJECT PORTFOLIO SUMMARY\n"]
        
        # Group by status
        by_status = self._group_projects_by_status(projects)
        
        for status, project_list in by_status.items():
            if project_list:
                summary_parts.append(f"\n{status.upper()} PROJECTS ({len(project_list)}):")
                for project in project_list:
                    project_line = f"- {project['name']}"
                    if project.get('number'):
                        project_line = f"- {project['number']}. {project['name']}"
                    if project.get('product_owner'):
                        project_line += f" (Owner: {project['product_owner']})"
                    summary_parts.append(project_line)
        
        summary_parts.append(f"\nTOTAL PROJECTS: {len(projects)}")
        
        return '\n'.join(summary_parts)

    def _group_projects_by_status(self, projects: List[Dict]) -> Dict[str, List[Dict]]:
        """Group projects by their status"""
        
        grouped = {
            "in_progress": [],
            "completed": [],
            "halted": []
        }
        
        for project in projects:
            status = project["status"].lower()
            if "in progress" in status:
                grouped["in_progress"].append(project)
            elif "completed" in status:
                grouped["completed"].append(project)
            elif "halted" in status:
                grouped["halted"].append(project)
        
        return grouped

    def _extract_project_section(self, content: str, project_name: str) -> str:
        """Extract specific project section from mixed content"""
        
        lines = content.split('\n')
        project_section = []
        in_target_section = False
        
        for line in lines:
            line_lower = line.lower()
            
            # Start of target project section
            if project_name in line_lower and any(indicator in line_lower for indicator in ["status:", "description:", "project"]):
                in_target_section = True
                project_section.append(line)
            # End of section (next project or major break)
            elif in_target_section and (
                line.startswith("**") or 
                "status:" in line_lower and project_name not in line_lower or
                line.strip().isdigit() and line.strip() + "." in content
            ):
                break
            # Continue collecting lines in section
            elif in_target_section:
                project_section.append(line)
        
        return '\n'.join(project_section) if project_section else content
   
    def _query_bm25_only(self, query: str, top_k: int) -> Dict[str, Any]:
        """Query using BM25 keyword search only with enhanced preprocessing"""
        try:
            if not self.bm25_index or not self.bm25_documents:
                return {"success": True, "documents": [], "message": "No BM25 index available"}
            
            # CHANGE this line:
            enhanced_query = self._enhance_query_for_retrieval(query) or query
            
            # Tokenize enhanced query
            query_tokens = enhanced_query.lower().split()
            
            # Get BM25 scores
            bm25_scores = self.bm25_index.get_scores(query_tokens)
            
            # Get top results
            top_indices = np.argsort(bm25_scores)[::-1][:top_k]
            
            formatted_results = []
            for idx in top_indices:
                if idx < len(self.bm25_documents) and bm25_scores[idx] > 0:
                    # Extract actual content for BM25 results too
                    actual_content = self._extract_actual_content(self.bm25_documents[idx])
                    
                    formatted_results.append({
                        "content": actual_content,
                        "full_content": self.bm25_documents[idx],
                        "metadata": self.bm25_metadata[idx],
                        "bm25_score": float(bm25_scores[idx]),
                        "similarity_score": float(bm25_scores[idx] / np.max(bm25_scores)) if max(bm25_scores) > 0 else 0,
                        "source_file": self.bm25_metadata[idx].get("source_file", "Unknown"),
                        "document_id": self.bm25_metadata[idx].get("document_id", "Unknown")
                    })
            
            # Prepare the result dictionary
            results = {
                "success": True,
                "query": query,
                "enhanced_query": enhanced_query,
                "documents": formatted_results,
                "total_results": len(formatted_results),
                "retrieval_method": "bm25_only"
            }
            
            # Add direct answer extraction if we have documents
            if results.get("success") and results.get("documents"):
                # Extract direct answer
                direct_answer = self.extract_direct_answer(query, results["documents"])
                results["direct_answer"] = direct_answer
            
            return results
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error in BM25 query: {e}")
            return {"success": False, "error": str(e)}
    
    def _query_hybrid(self, query: str, top_k: int) -> Dict[str, Any]:
        """Query using hybrid search (vector + BM25)"""
        try:
            # Get vector results
            vector_results = self._query_vector_only(query, top_k * 2)  # Get more for fusion
            
            # Get BM25 results
            bm25_results = self._query_bm25_only(query, top_k * 2)
            
            if not vector_results["success"] or not bm25_results["success"]:
                return vector_results if vector_results["success"] else bm25_results
            
            # Implement reciprocal rank fusion (RRF)
            fused_results = self._reciprocal_rank_fusion(
                vector_results["documents"], 
                bm25_results["documents"], 
                top_k
            )
            
            results = {
                "success": True,
                "query": query,
                "documents": fused_results,
                "total_results": len(fused_results),
                "retrieval_method": "hybrid_vector_bm25"
            }
            
            # Add direct answer extraction if we have documents
            if results.get("success") and results.get("documents"):
                # Extract direct answer
                direct_answer = self.extract_direct_answer(query, results["documents"])
                results["direct_answer"] = direct_answer
            
            return results
        
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error in hybrid query: {e}")
            return {"success": False, "error": str(e)}
    
    def _query_contextual_hybrid(self, query: str, top_k: int) -> Dict[str, Any]:
        """
        Query using Anthropic's Contextual Hybrid approach
        (Contextual Embeddings + Contextual BM25 + potential reranking)
        """
        try:
            # First, enhance query with context if needed
            enhanced_query = self._enhance_query_with_context(query)
            
            # Get hybrid results with enhanced query
            hybrid_results = self._query_hybrid(enhanced_query, top_k * 3)  # Get more for reranking
            
            if not hybrid_results["success"]:
                return hybrid_results
            
            # Apply reranking if we have enough results
            if len(hybrid_results["documents"]) > top_k:
                reranked_results = self._apply_contextual_reranking(
                    query, hybrid_results["documents"], top_k
                )
            else:
                reranked_results = hybrid_results["documents"][:top_k]
            
            results = {
                "success": True,
                "query": query,
                "enhanced_query": enhanced_query,
                "documents": reranked_results,
                "total_results": len(reranked_results),
                "retrieval_method": "contextual_hybrid",
                "contextual_enhancement": True
            }

            if results.get("success") and results.get("documents"):
                # Extract direct answer
                direct_answer = self.extract_direct_answer(query, results["documents"])
                results["direct_answer"] = direct_answer
            
            return results
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error in contextual hybrid query: {e}")
            return {"success": False, "error": str(e)}
    
    def _enhance_query_with_context(self, query: str) -> str:
        """Enhance query with contextual information"""
        # For now, return original query
        # This could be enhanced with query expansion techniques
        return query
    
    def _apply_contextual_reranking(self, query: str, documents: List[Dict[str, Any]], 
                                  top_k: int) -> List[Dict[str, Any]]:
        """Apply contextual reranking to improve result relevance"""
        try:
            # Simple relevance-based reranking
            # This could be enhanced with external reranking models (Cohere, etc.)
            
            reranked_docs = []
            for doc in documents:
                content = doc["content"].lower()
                query_lower = query.lower()
                
                # Calculate relevance score based on multiple factors
                relevance_score = 0
                
                # Exact phrase matches
                if query_lower in content:
                    relevance_score += 10
                
                # Word overlap
                query_words = set(query_lower.split())
                content_words = set(content.split())
                overlap = len(query_words.intersection(content_words))
                relevance_score += overlap * 2
                
                # Contextual metadata boost
                if doc["metadata"].get("has_context", False):
                    relevance_score += 1
                
                # Combine with original similarity score
                original_score = doc.get("similarity_score", 0)
                combined_score = (original_score * 0.7) + (relevance_score * 0.3)
                
                doc["rerank_score"] = combined_score
                reranked_docs.append(doc)
            
            # Sort by combined score
            reranked_docs.sort(key=lambda x: x.get("rerank_score", 0), reverse=True)
            
            return reranked_docs[:top_k]
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error in contextual reranking: {e}")
            return documents[:top_k]
    
    def _reciprocal_rank_fusion(self, vector_results: List[Dict[str, Any]], 
                              bm25_results: List[Dict[str, Any]], top_k: int) -> List[Dict[str, Any]]:
        """Implement Reciprocal Rank Fusion for combining results"""
        try:
            # Create a scoring map
            doc_scores = {}
            
            # Add vector scores (weight: 0.6)
            for i, doc in enumerate(vector_results):
                doc_id = doc["metadata"].get("chunk_id", f"vec_{i}")
                rrf_score = 1 / (60 + i + 1)  # RRF with k=60
                doc_scores[doc_id] = {
                    "document": doc,
                    "vector_score": rrf_score * 0.6,
                    "bm25_score": 0
                }
            
            # Add BM25 scores (weight: 0.4)
            for i, doc in enumerate(bm25_results):
                doc_id = doc["metadata"].get("chunk_id", f"bm25_{i}")
                rrf_score = 1 / (60 + i + 1)  # RRF with k=60
                
                if doc_id in doc_scores:
                    doc_scores[doc_id]["bm25_score"] = rrf_score * 0.4
                else:
                    doc_scores[doc_id] = {
                        "document": doc,
                        "vector_score": 0,
                        "bm25_score": rrf_score * 0.4
                    }
            
            # Calculate final scores and sort
            final_results = []
            for doc_id, scores in doc_scores.items():
                final_score = scores["vector_score"] + scores["bm25_score"]
                doc = scores["document"]
                doc["fusion_score"] = final_score
                doc["vector_contribution"] = scores["vector_score"]
                doc["bm25_contribution"] = scores["bm25_score"]
                final_results.append(doc)
            
            # Sort by fusion score and return top k
            final_results.sort(key=lambda x: x["fusion_score"], reverse=True)
            return final_results[:top_k]
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error in reciprocal rank fusion: {e}")
            return (vector_results + bm25_results)[:top_k]
    
    def _format_query_results(self, results: Dict[str, Any], query: str, method: str) -> Dict[str, Any]:
        """Format query results consistently"""
        try:
            if not results or not results.get("documents"):
                return {
                    "success": True,
                    "query": query,
                    "documents": [],
                    "message": "No relevant documents found",
                    "retrieval_method": method
                }
            
            # Handle ChromaDB response format - documents/metadatas/distances are nested lists
            documents = results["documents"][0] if results["documents"] and len(results["documents"]) > 0 else []
            metadatas = results["metadatas"][0] if results["metadatas"] and len(results["metadatas"]) > 0 else []
            distances = results["distances"][0] if results["distances"] and len(results["distances"]) > 0 else []
            
            # Ensure all lists have the same length
            min_length = min(len(documents), len(metadatas), len(distances)) if documents and metadatas and distances else 0
            
            formatted_results = []
            for i in range(min_length):
                doc = documents[i] if i < len(documents) else ""
                metadata = metadatas[i] if i < len(metadatas) else {}
                distance = distances[i] if i < len(distances) else 1.0
                
                # Fix similarity calculation - handle negative values
                if distance is not None:
                    # For very large distances, use exponential decay
                    if distance > 1.0:
                        similarity_score = 1.0 / (1.0 + distance)
                    else:
                        similarity_score = max(0, 1 - distance)
                else:
                    similarity_score = 0
                
                # Extract actual content from contextual content
                actual_content = self._extract_actual_content(doc)
                
                formatted_results.append({
                    "content": actual_content,       # Show actual content, not contextual
                    "full_content": doc,             # Keep full contextual content available
                    "metadata": metadata,
                    "similarity_score": similarity_score,
                    "distance": distance,
                    "source_file": metadata.get("source_file", "Unknown") if metadata else "Unknown",
                    "document_id": metadata.get("document_id", None) if metadata else None,
                    "title": metadata.get("title") if metadata else None,
                    "status": metadata.get("status") if metadata else None,
                    "has_context": metadata.get("has_context", False) if metadata else False,
                    "context_added": metadata.get("context_added", 0) if metadata else 0
                })
            
            return {
                "success": True,
                "query": query,
                "documents": formatted_results,
                "total_results": len(formatted_results),
                "retrieval_method": method
            }
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error formatting results: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_actual_content(self, contextual_content: str) -> str:
        """Extract the actual document content from contextual content"""
        try:
            # Split by double newline to separate context from actual content
            if "\n\n" in contextual_content:
                parts = contextual_content.split('\n\n', 1)
                if len(parts) > 1:
                    # The second part should be the actual content
                    return parts[1].strip()
            
            # Fallback: if no clear separation, return the content after "Context:"
            if contextual_content.startswith("Context:") or "Context:" in contextual_content[:100]:
                # Find the end of the context section
                lines = contextual_content.split('\n')
                content_lines = []
                found_content = False
                
                for line in lines:
                    if found_content:
                        content_lines.append(line)
                    elif line.strip() == "" and len(content_lines) == 0:
                        # Empty line after context, content starts next
                        found_content = True
                    elif not line.startswith("Context:") and not line.startswith("This") and len(line.strip()) > 0 and found_content == False:
                        # Likely start of actual content
                        found_content = True
                        content_lines.append(line)
                
                if content_lines:
                    return '\n'.join(content_lines).strip()
            
            # If no context markers found, return as-is
            return contextual_content.strip()
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error extracting actual content: {e}")
            return contextual_content
    
    def _update_file_versions(self, file_stem: str, doc_id: str, file_hash: str, processing_record: Dict[str, Any]):
        """Update file versions tracking"""
        if file_stem not in self.file_versions:
            self.file_versions[file_stem] = []
        
        self.file_versions[file_stem].append({
            "document_id": doc_id,
            "file_hash": file_hash,
            "processed_date": processing_record["processed_date"],
            "version_number": processing_record.get("version_number", 1),
            "is_latest": processing_record.get("is_latest_version", True),
            "contextual_enhanced": processing_record.get("contextual_enhanced", False)
        })
    
    def _update_processing_statistics(self, processed: int, skipped: int, overwritten: int, 
                                    errors: int, processing_time: float, contextual_enhancements: int):
        """Update processing statistics"""
        self.processing_stats = {
            "last_run": {
                "timestamp": datetime.now().isoformat(),
                "files_processed": processed,
                "files_skipped": skipped,
                "files_overwritten": overwritten,
                "files_with_errors": errors,
                "contextual_enhancements": contextual_enhancements,
                "processing_time_seconds": processing_time,
                "duplicate_mode": self.duplicate_mode.value,
                "chunking_strategy": self.chunking_strategy.value,
                "retrieval_strategy": self.retrieval_strategy.value
            },
            "totals": {
                "total_files_processed": self.processing_stats.get("totals", {}).get("total_files_processed", 0) + processed,
                "total_contextual_enhancements": self.processing_stats.get("totals", {}).get("total_contextual_enhancements", 0) + contextual_enhancements,
                "total_processing_runs": self.processing_stats.get("totals", {}).get("total_processing_runs", 0) + 1,
                "total_processing_time": self.processing_stats.get("totals", {}).get("total_processing_time", 0) + processing_time
            }
        }
    
    def _generate_document_id(self, file_path: str) -> str:
        """Generate unique document ID based on file path"""
        content = f"{file_path}_contextual_processed"
        return hashlib.md5(content.encode()).hexdigest()[:16]
    
    # Statistics and reporting methods
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get comprehensive processing statistics with contextual enhancement info"""
        try:
            total_files = len(self.processed_files_log)
            successful = sum(1 for record in self.processed_files_log.values() 
                           if record.get("status") == "success")
            failed = sum(1 for record in self.processed_files_log.values() 
                        if record.get("status") == "failed")
            contextual_enhanced = sum(1 for record in self.processed_files_log.values() 
                                    if record.get("contextual_enhanced", False))
            
            # Count versions
            total_file_groups = len(self.file_versions)
            total_versions = sum(len(versions) for versions in self.file_versions.values())
            
            # Get collection stats
            documents_count = 0
            if self.collection:
                try:
                    documents_count = self.collection.count()
                except:
                    pass
            
            # Calculate storage info
            total_size = sum(record.get("file_size", 0) for record in self.processed_files_log.values())
            
            stats = {
                "processing_summary": {
                    "total_files_processed": total_files,
                    "successful_processing": successful,
                    "failed_processing": failed,
                    "contextual_enhanced_files": contextual_enhanced,
                    "contextual_enhancement_rate": round((contextual_enhanced / total_files * 100) if total_files > 0 else 0, 2),
                    "success_rate": round((successful / total_files * 100) if total_files > 0 else 0, 2)
                },
                "version_control": {
                    "unique_file_groups": total_file_groups,
                    "total_versions": total_versions,
                    "average_versions_per_file": round(total_versions / total_file_groups, 2) if total_file_groups > 0 else 0
                },
                "storage_info": {
                    "vector_documents_count": documents_count,
                    "bm25_documents_count": len(self.bm25_documents),
                    "total_file_size_bytes": total_size,
                    "total_file_size_mb": round(total_size / (1024 * 1024), 2)
                },
                "anthropic_features": {
                    "contextual_embeddings": True,
                    "contextual_bm25": self.bm25_index is not None,
                    "hybrid_search": True,
                    "claude_enhancement": self.claude_client is not None,
                    "contextual_cache_size": len(self.contextual_cache)
                },
                "configuration": {
                    "duplicate_handling_mode": self.duplicate_mode.value,
                    "chunking_strategy": self.chunking_strategy.value,
                    "retrieval_strategy": self.retrieval_strategy.value,
                    "embedding_model": self.embedding_model_name,
                    "chunk_size": self.chunk_size,
                    "chunk_overlap": self.chunk_overlap,
                    "context_tokens": self.context_tokens,
                    "supported_extensions": list(self.supported_extensions)
                },
                "paths": {
                    "documents_directory": str(self.documents_dir),
                    "chromadb_path": str(self.chroma_db_path),
                    "collection_name": self.collection_name
                },
                "last_run": self.processing_stats.get("last_run", {}),
                "totals": self.processing_stats.get("totals", {}),
                "timestamp": datetime.now().isoformat()
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error getting processing stats: {e}")
            return {"error": str(e)}
    
    def get_processed_files_list(self) -> List[Dict[str, Any]]:
        """Get detailed list of processed files with contextual enhancement info"""
        try:
            files_list = []
            
            for filename, record in self.processed_files_log.items():
                file_info = {
                    "filename": filename,
                    "status": record.get("status"),
                    "processed_date": record.get("processed_date"),
                    "document_id": record.get("document_id"),
                    "chunks_created": record.get("chunks_created", 0),
                    "contextual_enhanced": record.get("contextual_enhanced", False),
                    "file_path": record.get("file_path"),
                    "file_hash": record.get("file_hash"),
                    "file_size": record.get("file_size"),
                    "file_size_mb": round(record.get("file_size", 0) / (1024 * 1024), 2),
                    "processing_action": record.get("processing_action", "unknown"),
                    "duplicate_mode": record.get("duplicate_mode", "unknown"),
                    "chunking_strategy": record.get("chunking_strategy", "unknown"),
                    "retrieval_strategy": record.get("retrieval_strategy", "unknown")
                }
                
                # Add version info if available
                if record.get("version_number"):
                    file_info["version_number"] = record["version_number"]
                    file_info["is_latest_version"] = record.get("is_latest_version", True)
                
                # Add file stem for version grouping
                file_info["file_stem"] = record.get("file_stem", filename.rsplit('.', 1)[0])
                
                files_list.append(file_info)
            
            return sorted(files_list, key=lambda x: x.get("processed_date", ""), reverse=True)
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error getting processed files list: {e}")
            return []
    
    def get_version_summary(self) -> Dict[str, Any]:
        """Get comprehensive version summary with contextual enhancement info"""
        try:
            version_summary = {}
            
            for file_stem, versions in self.file_versions.items():
                latest_version = max(versions, key=lambda v: v.get("processed_date", ""))
                contextual_enhanced_versions = sum(1 for v in versions if v.get("contextual_enhanced", False))
                
                version_summary[file_stem] = {
                    "total_versions": len(versions),
                    "contextual_enhanced_versions": contextual_enhanced_versions,
                    "latest_version": latest_version,
                    "all_versions": sorted(versions, key=lambda v: v.get("processed_date", ""), reverse=True),
                    "has_multiple_versions": len(versions) > 1,
                    "first_processed": min(versions, key=lambda v: v.get("processed_date", "")).get("processed_date", "Unknown"),
                    "last_processed": latest_version.get("processed_date", "Unknown")
                }
            
            return {
                "file_groups": version_summary,
                "summary": {
                    "total_file_groups": len(version_summary),
                    "files_with_multiple_versions": sum(1 for info in version_summary.values() 
                                                       if info["has_multiple_versions"]),
                    "total_versions": sum(info["total_versions"] for info in version_summary.values()),
                    "total_contextual_enhanced": sum(info["contextual_enhanced_versions"] for info in version_summary.values()),
                    "duplicate_handling_mode": self.duplicate_mode.value,
                    "retrieval_strategy": self.retrieval_strategy.value
                },
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error getting version summary: {e}")
            return {"error": str(e)}
    
    # Utility methods
    def set_duplicate_mode(self, mode: DuplicateHandlingMode):
        """Change duplicate handling mode"""
        self.duplicate_mode = mode
        logger.info(f"Ã°Å¸â€â€ž Duplicate handling mode changed to: {mode.value}")
    
    def set_chunking_strategy(self, strategy: ChunkingStrategy):
        """Change chunking strategy"""
        self.chunking_strategy = strategy
        logger.info(f"Ã°Å¸â€â€ž Chunking strategy changed to: {strategy.value}")
    
    def set_retrieval_strategy(self, strategy: RetrievalStrategy):
        """Change retrieval strategy"""
        self.retrieval_strategy = strategy
        logger.info(f"Ã°Å¸â€â€ž Retrieval strategy changed to: {strategy.value}")
    
    def cleanup_old_versions(self, keep_latest_n: int = 2) -> Dict[str, Any]:
        """Clean up old versions to optimize storage"""
        try:
            removed_count = 0
            
            for file_stem, versions in self.file_versions.items():
                if len(versions) <= keep_latest_n:
                    continue
                
                # Sort versions by processed date (newest first)
                sorted_versions = sorted(
                    versions, 
                    key=lambda v: v.get("processed_date", ""), 
                    reverse=True
                )
                
                # Keep only the latest N versions
                versions_to_keep = sorted_versions[:keep_latest_n]
                versions_to_remove = sorted_versions[keep_latest_n:]
                
                # Remove old versions from ChromaDB and BM25
                for version in versions_to_remove:
                    doc_id = version.get("document_id")
                    if doc_id and self.collection:
                        try:
                            # Remove from ChromaDB
                            results = self.collection.get(
                                where={"document_id": doc_id},
                                include=["documents"]
                            )
                            if results and results.get("ids"):
                                self.collection.delete(ids=results["ids"])
                                removed_count += len(results["ids"])
                            
                            # Remove from BM25
                            indices_to_remove = []
                            for i, metadata in enumerate(self.bm25_metadata):
                                if metadata.get("document_id") == doc_id:
                                    indices_to_remove.append(i)
                            
                            for i in sorted(indices_to_remove, reverse=True):
                                if i < len(self.bm25_documents):
                                    del self.bm25_documents[i]
                                    del self.bm25_metadata[i]
                                    
                        except Exception as e:
                            logger.warning(f"Ã¢Å¡ Ã¯Â¸Â Could not remove chunks for {doc_id}: {e}")
                
                # Update file versions tracking
                self.file_versions[file_stem] = versions_to_keep
            
            # Rebuild BM25 index
            if self.bm25_documents:
                tokenized_docs = [doc.lower().split() for doc in self.bm25_documents]
                self.bm25_index = BM25Okapi(tokenized_docs)
            
            # Clean up processing log
            valid_doc_ids = set()
            for versions in self.file_versions.values():
                for version in versions:
                    valid_doc_ids.add(version.get("document_id"))
            
            keys_to_remove = []
            for key, processed_info in self.processed_files_log.items():
                if processed_info.get("document_id") not in valid_doc_ids:
                    keys_to_remove.append(key)
            
            for key in keys_to_remove:
                del self.processed_files_log[key]
            
            # Save updated logs
            self._save_processing_logs()
            
            logger.info(f"Ã°Å¸Â§Â¹ Cleanup complete: removed {removed_count} chunks from old versions")
            
            return {
                "success": True,
                "chunks_removed": removed_count,
                "files_cleaned": len([f for f, v in self.file_versions.items() if len(v) > keep_latest_n]),
                "kept_versions_per_file": keep_latest_n
            }
            
        except Exception as e:
            logger.error(f"Ã¢ÂÅ’ Error during cleanup: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
# Terminal interface for manual execution
def print_statistics(rag_service: AnthropicContextualRAGService):
    """Print processing statistics in terminal-friendly format"""
    print("\n" + "="*80)
    print("Ã°Å¸â€œÅ  ANTHROPIC CONTEXTUAL RAG STATISTICS")
    print("="*80)
    
    stats = rag_service.get_processing_statistics()
    
    if "error" in stats:
        print(f"Ã¢ÂÅ’ Error: {stats['error']}")
        return
    
    # Processing Summary
    summary = stats.get("processing_summary", {})
    print(f"\nÃ°Å¸â€œË† Processing Summary:")
    print(f"   Total Files Processed: {summary.get('total_files_processed', 0)}")
    print(f"   Successful: {summary.get('successful_processing', 0)}")
    print(f"   Failed: {summary.get('failed_processing', 0)}")
    print(f"   Contextually Enhanced: {summary.get('contextual_enhanced_files', 0)}")
    print(f"   Enhancement Rate: {summary.get('contextual_enhancement_rate', 0)}%")
    print(f"   Success Rate: {summary.get('success_rate', 0)}%")
    
    # Anthropic Features
    anthropic_features = stats.get("anthropic_features", {})
    print(f"\nÃ°Å¸Â§  Anthropic Features:")
    print(f"   Contextual Embeddings: {'Ã¢Å“â€¦' if anthropic_features.get('contextual_embeddings') else 'Ã¢ÂÅ’'}")
    print(f"   Contextual BM25: {'Ã¢Å“â€¦' if anthropic_features.get('contextual_bm25') else 'Ã¢ÂÅ’'}")
    print(f"   Hybrid Search: {'Ã¢Å“â€¦' if anthropic_features.get('hybrid_search') else 'Ã¢ÂÅ’'}")
    print(f"   Claude Enhancement: {'Ã¢Å“â€¦' if anthropic_features.get('claude_enhancement') else 'Ã¢ÂÅ’'}")
    print(f"   Contextual Cache Size: {anthropic_features.get('contextual_cache_size', 0)}")
    
    # Storage Info
    storage = stats.get("storage_info", {})
    print(f"\nÃ°Å¸â€™Â¾ Storage Information:")
    print(f"   Vector Documents: {storage.get('vector_documents_count', 0)}")
    print(f"   BM25 Documents: {storage.get('bm25_documents_count', 0)}")
    print(f"   Total File Size: {storage.get('total_file_size_mb', 0)} MB")
    
    # Configuration
    config = stats.get("configuration", {})
    print(f"\nÃ¢Å¡â„¢Ã¯Â¸Â Configuration:")
    print(f"   Retrieval Strategy: {config.get('retrieval_strategy', 'Unknown')}")
    print(f"   Chunking Strategy: {config.get('chunking_strategy', 'Unknown')}")
    print(f"   Embedding Model: {config.get('embedding_model', 'Unknown')}")
    print(f"   Chunk Size: {config.get('chunk_size', 0)}")
    print(f"   Context Tokens: {config.get('context_tokens', 0)}")

def print_processed_files(rag_service: AnthropicContextualRAGService):
    """Print processed files list in terminal-friendly format"""
    print("\n" + "="*80)
    print("Ã°Å¸â€œÂ PROCESSED FILES ")
    print("="*80)
    
    files = rag_service.get_processed_files_list()
    
    if not files:
        print("No files have been processed yet.")
        return
    
    print(f"\nTotal files: {len(files)}")
    print("-" * 80)
    
    for i, file_info in enumerate(files, 1):
        status_icon = "Ã¢Å“â€¦" if file_info.get("status") == "success" else "Ã¢ÂÅ’"
        context_icon = "Ã°Å¸Â§ " if file_info.get("contextual_enhanced") else "Ã°Å¸â€œâ€ž"
        
        print(f"\n{i}. {status_icon} {context_icon} {file_info.get('filename', 'Unknown')}")
        print(f"   Status: {file_info.get('status', 'Unknown')}")
        print(f"   Size: {file_info.get('file_size_mb', 0)} MB")
        print(f"   Chunks: {file_info.get('chunks_created', 0)}")
        print(f"   Contextual Enhanced: {'Yes' if file_info.get('contextual_enhanced') else 'No'}")
        print(f"   Processed: {file_info.get('processed_date', 'Unknown')}")
        print(f"   Retrieval Strategy: {file_info.get('retrieval_strategy', 'Unknown')}")
        
        if file_info.get("version_number"):
            print(f"   Version: {file_info.get('version_number')}")

def print_version_summary(rag_service: AnthropicContextualRAGService):
    """Print version summary in terminal-friendly format"""
    print("\n" + "="*80)
    print("Ã°Å¸â€œâ€¹ VERSION SUMMARY (ANTHROPIC ENHANCED)")
    print("="*80)
    
    version_info = rag_service.get_version_summary()
    
    if "error" in version_info:
        print(f"Ã¢ÂÅ’ Error: {version_info['error']}")
        return
    
    summary = version_info.get("summary", {})
    print(f"\nÃ°Å¸â€œÅ  Summary:")
    print(f"   Total File Groups: {summary.get('total_file_groups', 0)}")
    print(f"   Files with Multiple Versions: {summary.get('files_with_multiple_versions', 0)}")
    print(f"   Total Versions: {summary.get('total_versions', 0)}")
    print(f"   Contextually Enhanced: {summary.get('total_contextual_enhanced', 0)}")
    print(f"   Retrieval Strategy: {summary.get('retrieval_strategy', 'Unknown')}")
    
    file_groups = version_info.get("file_groups", {})
    if file_groups:
        print(f"\nÃ°Å¸â€œÂ File Groups:")
        print("-" * 80)
        
        for file_stem, info in file_groups.items():
            versions_icon = "Ã°Å¸â€œÅ¡" if info.get("has_multiple_versions") else "Ã°Å¸â€œâ€ž"
            enhanced_count = info.get("contextual_enhanced_versions", 0)
            
            print(f"\n{versions_icon} {file_stem}")
            print(f"   Total Versions: {info.get('total_versions', 0)}")
            print(f"   Contextually Enhanced: {enhanced_count}")
            print(f"   First Processed: {info.get('first_processed', 'Unknown')}")
            print(f"   Last Processed: {info.get('last_processed', 'Unknown')}")

# Update the main function - find this section and replace it:

def main():
    """Main function for terminal execution"""
    parser = argparse.ArgumentParser(description="Anthropic Contextual RAG Service - Advanced Document Processing")
    parser.add_argument("--docs-dir", help="Documents directory path")
    parser.add_argument("--db-path", default="./chroma_db", help="ChromaDB path")
    parser.add_argument("--duplicate-mode", choices=["skip", "overwrite", "version", "timestamp"], 
                       default="skip", help="Duplicate handling mode")
    parser.add_argument("--chunking", choices=["fixed_size", "semantic", "sentence", "paragraph"],
                       default="semantic", help="Chunking strategy")
    parser.add_argument("--retrieval", choices=["vector_only", "bm25_only", "hybrid", "contextual_hybrid"],
                       default="contextual_hybrid", help="Retrieval strategy")
    parser.add_argument("--process", action="store_true", help="Process all documents")
    parser.add_argument("--stats", action="store_true", help="Show processing statistics")
    parser.add_argument("--files", action="store_true", help="Show processed files list")
    parser.add_argument("--versions", action="store_true", help="Show version summary")
    parser.add_argument("--query", help="Query the vector database")
    parser.add_argument("--cleanup", type=int, metavar="N", help="Cleanup old versions, keep latest N")
    parser.add_argument("--check-cache", action="store_true", help="Check model cache status")
    parser.add_argument("--clear-cache", action="store_true", help="Clear model cache")
    parser.add_argument("--quiet", "-q", action="store_true", help="Suppress initialization output")
    parser.add_argument("--verbose", "-v", action="store_true", help="Show detailed initialization output")
    
    args = parser.parse_args()
    
    # Convert string arguments to enums
    duplicate_mode = DuplicateHandlingMode(args.duplicate_mode)
    chunking_strategy = ChunkingStrategy(args.chunking)
    retrieval_strategy = RetrievalStrategy(args.retrieval)
    
    # Determine if we should use quiet mode
    # Use quiet mode for info commands unless --verbose is explicitly used
    use_quiet_mode = False
    if args.verbose:
        use_quiet_mode = False  
    elif args.quiet:
        use_quiet_mode = True   
    else:
        # Auto-quiet for these commands unless verbose is specified
        use_quiet_mode = any([args.stats, args.files, args.versions, args.query, args.check_cache, args.clear_cache])
    
    try:
        # Initialize Anthropic RAG service
        if not use_quiet_mode:
            print("Ã°Å¸Å¡â‚¬ Initializing Anthropic Contextual RAG Service...")
            
        rag_service = AnthropicContextualRAGService(
            documents_directory=args.docs_dir,
            chroma_db_path=args.db_path,
            duplicate_mode=duplicate_mode,
            chunking_strategy=chunking_strategy,
            retrieval_strategy=retrieval_strategy,
            quiet_mode=use_quiet_mode
        )
        
        if not use_quiet_mode:
            print("Ã¢Å“â€¦ Anthropic RAG Service initialized successfully!")
        
        # Execute requested operations
        if args.check_cache:
            if not use_quiet_mode:
                print("\nÃ°Å¸â€Â Checking model cache status...")
            cache_path = rag_service._get_model_cache_path()
            is_cached = rag_service._check_model_cache(cache_path)
            print(f"Ã°Å¸â€œÂ Cache path: {cache_path}")
            print(f"Ã¢Å“â€¦ Is cached: {'Yes' if is_cached else 'No'}")

        if args.clear_cache:
            if not use_quiet_mode:
                print("\nÃ°Å¸â€”â€˜Ã¯Â¸Â Clearing model cache...")
            result = rag_service._clear_model_cache()
            print(f"Ã¢Å“â€¦ Cache cleared: {'Success' if result else 'Failed'}")
        
        if args.process:
            print("\nÃ°Å¸Â§  Processing all documents with Anthropic's Contextual Retrieval...")
            result = rag_service.process_all_documents()
            if result["success"]:
                print(f"Ã¢Å“â€¦ Processing completed successfully!")
                print(f"   Files processed: {result['files_processed']}")
                print(f"   Files skipped: {result['files_skipped']}")
                print(f"   Files overwritten: {result['files_overwritten']}")
                print(f"   Contextual enhancements: {result.get('contextual_enhancements', 0)}")
                print(f"   Processing time: {result['processing_time']} seconds")
            else:
                print(f"Ã¢ÂÅ’ Processing failed: {result.get('error', 'Unknown error')}")
        
        if args.stats:
            print_statistics(rag_service)
        
        if args.files:
            print_processed_files(rag_service)
        
        if args.versions:
            print_version_summary(rag_service)
        
        if args.query:
            print(f"Ã°Å¸â€Â Querying with {retrieval_strategy.value}: '{args.query}'\n")
            results = rag_service.query_documents(args.query)
            if results["success"]:
                # Show direct answer first
                if results.get("direct_answer"):
                    print("Ã°Å¸Å½Â¯ DIRECT ANSWER:")
                    print(f"   {results['direct_answer']}")
                    print("-" * 80)
                
                if results["documents"]:
                    print(f"\nFound {len(results['documents'])} supporting documents:")
                    print("-" * 80)
                    for i, doc in enumerate(results["documents"], 1):
                        context_indicator = "Ã°Å¸Â§ " if doc.get("has_context") else "Ã°Å¸â€œâ€ž"
                        print(f"\n{i}. {context_indicator} Source: {doc['source_file']}")
                        print(f"   Similarity: {doc['similarity_score']:.3f}")
                        if doc.get("fusion_score"):
                            print(f"   Fusion Score: {doc['fusion_score']:.3f}")
                        
                        # Show content preview
                        actual_content = doc['content']
                        display_content = actual_content[:200] + "..." if len(actual_content) > 200 else actual_content
                        print(f"   Content: {display_content}")
                        print("-" * 40)
                else:
                    print("No supporting documents found.")
            else:
                print(f"Ã¢ÂÅ’ Query failed: {results.get('error', 'Unknown error')}")
        
        if args.cleanup:
            print(f"\nÃ°Å¸Â§Â¹ Cleaning up old versions, keeping latest {args.cleanup}...")
            result = rag_service.cleanup_old_versions(keep_latest_n=args.cleanup)
            if result["success"]:
                print(f"Ã¢Å“â€¦ Cleanup completed!")
                print(f"   Chunks removed: {result['chunks_removed']}")
                print(f"   Files cleaned: {result['files_cleaned']}")
            else:
                print(f"Ã¢ÂÅ’ Cleanup failed: {result.get('error', 'Unknown error')}")
        
    except Exception as e:
        print(f"\nÃ¢ÂÅ’ Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()

# Convenience functions for creating service instances
def create_anthropic_rag_service(
    documents_directory: Optional[str] = None,
    duplicate_mode: DuplicateHandlingMode = DuplicateHandlingMode.SKIP,
    retrieval_strategy: RetrievalStrategy = RetrievalStrategy.CONTEXTUAL_HYBRID
) -> AnthropicContextualRAGService:
    """
    Create Anthropic Contextual RAG Service
    
    Args:
        documents_directory: Custom documents directory path
        duplicate_mode: How to handle duplicate files
        retrieval_strategy: Retrieval method to use
        
    Returns:
        Initialized AnthropicContextualRAGService instance
    """
    return AnthropicContextualRAGService(
        documents_directory=documents_directory,
        duplicate_mode=duplicate_mode,
        retrieval_strategy=retrieval_strategy
    )

# Backward compatibility alias
EnhancedRAGService = AnthropicContextualRAGService

